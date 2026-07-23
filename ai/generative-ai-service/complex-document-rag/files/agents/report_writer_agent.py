from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import os
import uuid
import logging
import datetime
import math
import re
from typing import List, Optional
from docx.oxml.shared import OxmlElement
from docx.text.run import Run
import matplotlib.patches as mpatches
from matplotlib.table import Table
import numpy as np
import textwrap

from langchain_core.messages import HumanMessage
from contracts import Chunk, SectionDraft, ReportResult
from progress_bus import progress_bus

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

os.makedirs("charts", exist_ok=True)



_MD_TOKEN_RE = re.compile(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)')

def add_inline_markdown_paragraph(doc, text: str, style: str | None = None):
    """
    Creates a paragraph and renders lightweight inline Markdown:
      **bold** or __bold__ → bold run
      *italic* or _italic_ → italic run
    Everything else is plain text. No links/lists/code handling.

    style: optional python-docx paragraph style name (e.g. "List Bullet").
    """
    try:
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except KeyError:
        # Style missing from the template — fall back to a plain paragraph
        p = doc.add_paragraph()
    i = 0
    for m in _MD_TOKEN_RE.finditer(text):
        # leading text
        if m.start() > i:
            p.add_run(text[i:m.start()])
        token = m.group(0)
        # strip the markers
        if token.startswith('**') or token.startswith('__'):
            content = token[2:-2]
            run = p.add_run(content)
            run.bold = True
        else:
            content = token[1:-1]
            run = p.add_run(content)
            run.italic = True
        i = m.end()
    # trailing text
    if i < len(text):
        p.add_run(text[i:])
    return p

def add_table(doc, table_data):
    """Create a Word table from list of dicts or list of lists, robustly.

    Uses compact font sizing (8 pt body, 8.5 pt bold headers) with Oracle
    brand colours to avoid ugly line-wrapping and keep a corporate look.
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    if not table_data:
        return

    headers = []
    rows_normalized = []

    # Case 1: list of dicts
    if isinstance(table_data[0], dict):
        seen = set()
        for row in table_data:
            for k in row.keys():
                if k not in seen:
                    headers.append(k)
                    seen.add(k)
        rows_normalized = table_data

    # Case 2: list of lists
    elif isinstance(table_data[0], (list, tuple)):
        max_len = max(len(row) for row in table_data)
        headers = [f"Col {i+1}" for i in range(max_len)]
        for row in table_data:
            rows_normalized.append({headers[i]: row[i] if i < len(row) else ""
                                    for i in range(max_len)})
    else:
        headers = ["Value"]
        rows_normalized = [{"Value": str(row)} for row in table_data]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Allow autofit so Word can shrink columns to fit content
    table.autofit = True

    # ---- Helper: style a single cell ----
    _ORACLE_RED = RGBColor(0xC7, 0x46, 0x34)
    _WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    _CHARCOAL = RGBColor(0x31, 0x2D, 0x2A)
    _LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)

    def _shade_cell(cell, color_hex: str):
        """Apply background shading to a table cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def _style_cell(cell, text: str, *, bold: bool = False, font_size=Pt(8),
                    font_color=_CHARCOAL):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(text))
        run.font.size = font_size
        run.font.name = "Calibri"
        run.font.color.rgb = font_color
        run.bold = bold

    # ---- Header row (Oracle Red background, white text) ----
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        _style_cell(cell, str(h), bold=True, font_size=Pt(8.5), font_color=_WHITE)
        _shade_cell(cell, "C74634")

    # ---- Data rows (alternating white / light grey) ----
    for row_idx, row in enumerate(rows_normalized):
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            cell = row_cells[i]
            _style_cell(cell, str(row.get(h, "")))
            if row_idx % 2 == 1:
                _shade_cell(cell, "F5F5F5")


_SCALE_WORDS = (
    ("trillion", 1e12), ("tn", 1e12),
    ("billion", 1e9), ("bn", 1e9),
    ("million", 1e6),
    ("thousand", 1e3),
)


def _comparable_quantity(raw) -> tuple[float, str] | None:
    """
    Parse a cell into (magnitude, unit) for a direction-of-comparison check.

    Unlike _parse_numeric this keeps the unit and applies scale words, so
    "£600 million" and "£3.5 billion" become comparable. Returns None when the
    cell is not a clean single quantity — callers must treat that as "unknown"
    rather than guessing.

    Negative values are rejected on purpose: a cell like "-93%" is a *reduction*,
    where the numerically smaller value is the larger achievement. Annotating
    those by raw magnitude would invert the very comparison this guards.
    """
    if raw is None:
        return None
    s = str(raw).strip().lower()
    if s.upper() in ("N/A", "", "-", "—", "NONE"):
        return None
    if "-" in s.lstrip("-") or s.startswith("-") or "−" in s:
        return None
    s = s.replace(",", "")  # thousands separators would read as extra figures
    # A range ("50-80%") or a cell packing two figures has no single magnitude
    if len(re.findall(r"\d+(?:\.\d+)?", s)) != 1:
        return None

    m = re.search(r"(\d+(?:\.\d+)?)", s)
    if not m:
        return None
    val = float(m.group(1))

    for word, mult in _SCALE_WORDS:
        if re.search(rf"\b{word}\b", s) or s.rstrip("%").endswith(word):
            val *= mult
            break

    if "%" in s:
        unit = "pct"
    elif "£" in s:
        unit = "gbp"
    elif "$" in s:
        unit = "usd"
    elif re.search(r"\bmw\b", s):
        unit = "mw"
    elif re.search(r"\b(m|million|bn|billion)\b", s):
        unit = "count"
    else:
        unit = "num"
    return val, unit


def _higher_entity(row: dict, entities: list) -> str | None:
    """Which entity holds the numerically larger value, or None if undecidable."""
    if len(entities) != 2:
        return None
    a, b = entities
    qa, qb = _comparable_quantity(row.get(a)), _comparable_quantity(row.get(b))
    if not qa or not qb or qa[1] != qb[1] or qa[0] == qb[0]:
        return None
    return a if qa[0] > qb[0] else b


def _parse_numeric(raw) -> float | None:
    """Try to parse a table cell into a float, stripping currency/unit noise.

    Rejects cells that are predominantly text (e.g. credit ratings like "Aa3",
    qualitative labels like "Strong") by checking the ratio of digit characters
    to the overall string length.
    """
    if raw is None:
        return None
    s = str(raw).strip()
    if s.upper() in ("N/A", "", "-", "—", "NONE"):
        return None
    # Reject strings that are predominantly alphabetic (ratings, labels)
    digit_count = sum(1 for c in s if c.isdigit())
    if digit_count == 0:
        return None
    alpha_count = sum(1 for c in s if c.isalpha())
    # If more alphabetic chars than digits, likely a label not a number
    # (e.g. "Aa3" has 2 alpha, 1 digit → skip; "$14.2B" has 1 alpha, 3 digits → keep)
    if alpha_count > digit_count:
        return None
    cleaned = re.sub(r"[^\d.\-eE+]", "", s.replace(",", ""))
    if not cleaned or cleaned in (".", "-", "+"):
        return None
    try:
        val = float(cleaned)
        if math.isnan(val) or math.isinf(val):
            return None
        return val
    except (ValueError, TypeError):
        return None


def _chart_data_from_table(table: list[dict], entities: list[str]) -> dict:
    """Build **nested** chart_data from a validated comparison table.

    Returns ``{"Metric": {"Entity1": val, "Entity2": val}, ...}`` so that
    ``make_chart`` can render a proper grouped bar chart with one color per
    entity.  Rows where at least one entity has a numeric value are included
    so that partial data still produces a chart.
    """
    chart: dict[str, dict[str, float]] = {}
    for row in table:
        metric = str(row.get("Metric", ""))
        if not metric or metric == "Unknown Metric":
            continue
        entity_vals: dict[str, float] = {}
        for ent in entities:
            val = _parse_numeric(row.get(ent))
            if val is not None:
                entity_vals[ent] = val
        if entity_vals:
            chart[metric] = entity_vals
    if chart:
        logger.info(f"📊 _chart_data_from_table: extracted {len(chart)} chartable metrics from table")
    return chart


def _mine_chart_from_table(table: list[dict]) -> dict:
    """Generic fallback: extract flat {metric: value} from any table shape.

    Scans each row for a 'Metric' (or first string) key and picks the first
    parseable numeric value from the remaining columns.  Returns a flat dict
    suitable for ``make_chart`` in non-grouped mode.
    """
    chart: dict[str, float] = {}
    skip_cols = {"Metric", "Analysis", "Best Value", "Ranking", "analysis", "metric"}
    for row in table:
        if not isinstance(row, dict):
            continue
        metric = str(row.get("Metric", ""))
        if not metric:
            # Try the first string-valued key as metric name
            for k, v in row.items():
                if isinstance(v, str) and not _parse_numeric(v):
                    metric = v
                    break
        if not metric or metric == "Unknown Metric":
            continue
        # Find first numeric value
        for k, v in row.items():
            if k in skip_cols or str(k) == metric:
                continue
            val = _parse_numeric(v)
            if val is not None:
                chart[metric] = val
                break
    if chart:
        logger.info(f"📊 _mine_chart_from_table: mined {len(chart)} metrics from table")
    return chart


def _color_for_label(label: str, entities: list[str] | tuple[str, ...] | None,
                     base="#BFBFBF", e1="#C74634", e2="#312D2A") -> str:
    """Pick a bar color based on whether a label mentions one of the entities."""
    if not entities:
        return base
    lbl = label.lower()
    ents = [e for e in entities if isinstance(e, str)]
    if len(ents) >= 1 and ents[0].lower() in lbl:
        return e1
    if len(ents) >= 2 and ents[1].lower() in lbl:
        return e2
    return base


def detect_units(chart_data: dict, title: str = "") -> str:
    """Detect units of measure from chart data and title."""
    # Common patterns for currency
    currency_patterns = [
        (r'\$|USD|usd|dollar', 'USD'),
        (r'€|EUR|eur|euro', 'EUR'),
        (r'£|GBP|gbp|pound', 'GBP'),
        (r'¥|JPY|jpy|yen', 'JPY'),
        (r'₹|INR|inr|rupee', 'INR'),
    ]
    
    # Common patterns for other units - order matters!
    unit_patterns = [
        (r'million|millions|mn|mln|\$m|\$M', 'Million'),
        (r'billion|billions|bn|bln|\$b|\$B', 'Billion'),
        (r'thousand|thousands|k|\$k', 'Thousand'),
        (r'percentage|percent|%', '%'),
        (r'tonnes|tons|tonne|ton', 'Tonnes'),
        (r'co2e|CO2e|co2|CO2', 'CO2e'),
        (r'kwh|kWh|KWH', 'kWh'),
        (r'mwh|MWh|MWH', 'MWh'),
        (r'kg|kilogram|kilograms', 'kg'),
        (r'employees|headcount|people', 'Employees'),
        (r'days|day', 'Days'),
        (r'hours|hour|hrs', 'Hours'),
        (r'years|year|yrs', 'Years'),
    ]
    
    # Check title and keys for units - also check values if they're strings
    combined_text = title.lower() + " " + " ".join(str(k).lower() for k in chart_data.keys())
    # Also check string values which might contain unit info
    for v in chart_data.values():
        if isinstance(v, str):
            combined_text += " " + v.lower()
    
    detected_currency = None
    detected_scale = None
    detected_unit = None
    
    # Check for currency
    for pattern, unit in currency_patterns:
        if re.search(pattern, combined_text, re.IGNORECASE):
            detected_currency = unit
            break
    
    # Check for scale (million, billion, etc.)
    for pattern, unit in unit_patterns[:4]:  # First 4 are scales
        if re.search(pattern, combined_text, re.IGNORECASE):
            detected_scale = unit
            break
    
    # Check for other units
    for pattern, unit in unit_patterns[4:]:  # Rest are units
        if re.search(pattern, combined_text, re.IGNORECASE):
            detected_unit = unit
            break
    
    # Combine detected elements
    if detected_currency and detected_scale:
        return f"{detected_scale} {detected_currency}"
    elif detected_currency:
        # If we detect currency but no scale, look for financial context clues
        if 'revenue' in combined_text or 'sales' in combined_text or 'income' in combined_text:
            # Financial data without explicit scale often means millions
            if 'fy' in combined_text or 'fiscal' in combined_text or 'quarterly' in combined_text:
                return "Million USD"  # Corporate financials are typically in millions
            return detected_currency
        return detected_currency
    elif detected_unit:
        if detected_scale and detected_unit not in ['%', 'Employees', 'Days', 'Hours', 'Years']:
            return f"{detected_scale} {detected_unit}"
        return detected_unit
    elif detected_scale:
        # If we only have scale (like "Million") without currency, check for financial context
        if any(term in combined_text for term in ['revenue', 'cost', 'profit', 'income', 'sales', 'expense', 'financial']):
            return f"{detected_scale} USD"
        return detected_scale
    
    # For financial metrics without explicit units, default to "Million USD"
    if any(term in combined_text for term in ['revenue', 'sales', 'profit', 'income', 'cost', 'expense', 'financial', 'fiscal', 'fy20']):
        return "Million USD"
    
    return "Value"  # Default fallback


def format_value_with_units(value: float, units: str) -> str:
    """Format a value with appropriate precision based on units."""
    if '%' in units:
        return f"{value:.1f}%"
    elif 'Million' in units or 'Billion' in units:
        return f"{value:,.1f}"
    elif value >= 1000:
        return f"{value:,.0f}"
    else:
        return f"{value:.1f}"


def _make_grouped_chart(grouped_data: dict[str, dict[str, float]],
                        title: str, units: str,
                        entity_colors: dict[str, str]) -> str | None:
    """Render a grouped horizontal bar chart with one color per entity.

    ``grouped_data`` is ``{"Metric": {"Entity1": val, "Entity2": val}}``.
    ``entity_colors`` maps entity names to hex colours.
    Returns the path to the saved PNG or *None*.
    """
    import textwrap

    metrics = list(grouped_data.keys())
    if not metrics:
        return None

    all_entities = list(entity_colors.keys())
    n_entities = len(all_entities)
    bar_height = 0.8 / n_entities

    fig, ax = plt.subplots(figsize=(12, max(6, len(metrics) * 1.2)))

    for i, ent in enumerate(all_entities):
        positions = [m + i * bar_height for m in range(len(metrics))]
        vals = [grouped_data[metric].get(ent, 0) for metric in metrics]
        bars = ax.barh(positions, vals, height=bar_height,
                       label=ent, color=entity_colors[ent])
        for bar in bars:
            width = bar.get_width()
            if width != 0:
                formatted = format_value_with_units(width, units)
                ax.annotate(formatted,
                            xy=(width, bar.get_y() + bar.get_height() / 2),
                            xytext=(5, 0), textcoords="offset points",
                            ha="left", va="center", fontsize=8)

    center_offsets = [m + bar_height * (n_entities - 1) / 2 for m in range(len(metrics))]
    wrapped = ["\n".join(textwrap.wrap(m, width=35)) for m in metrics]
    ax.set_yticks(center_offsets)
    ax.set_yticklabels(wrapped)
    ax.set_xlabel(units)
    ax.set_title(title[:100])
    ax.legend(loc="lower right", framealpha=0.9)
    ax.grid(axis="x", linestyle="--", alpha=0.6)
    fig.tight_layout()

    filename = f"chart_{uuid.uuid4().hex}.png"
    path = os.path.join("charts", filename)
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


# Oracle-aligned palette for up to 10 entities in grouped charts
_ENTITY_PALETTE = [
    "#C74634",  # Oracle Red
    "#312D2A",  # Charcoal
    "#3A7CA5",  # Steel Blue
    "#747474",  # Mid Grey
    "#D4892A",  # Amber
    "#5B8C5A",  # Sage Green
    "#8E6FAD",  # Purple
    "#2E86AB",  # Ocean Blue
    "#A23B72",  # Magenta
    "#6A7B8B",  # Slate
]


def make_chart(chart_data: dict, title: str = "",
               entities: list[str] | tuple[str, ...] | None = None,
               units: str | None = None) -> str | None:
    """Generate a bar chart.

    * If any value in *chart_data* is itself a ``dict`` (nested / grouped
      data), a **grouped bar chart** is rendered with one colour per entity.
    * Otherwise a flat bar chart is rendered, with per-entity colour hints
      when *entities* is provided.

    Units are detected automatically or can be passed explicitly.
    """
    import textwrap

    if not chart_data:
        logger.info(f"📊 make_chart: empty chart_data for '{title}' — skipping")
        return None

    os.makedirs("charts", exist_ok=True)

    # Detect units early (before we potentially strip dict values)
    if not units:
        units = detect_units(chart_data, title)
    if units and units != "Value" and units.lower() not in title.lower():
        title = f"{title} ({units})"

    # --- Grouped mode: values are dicts keyed by entity -----------------------
    grouped: dict[str, dict[str, float]] = {}
    for k, v in chart_data.items():
        if isinstance(v, dict):
            nums = {}
            for sub_k, sub_v in v.items():
                try:
                    n = float(sub_v)
                    if not math.isnan(n) and not math.isinf(n):
                        nums[str(sub_k)] = n
                except Exception:
                    continue
            if nums:
                grouped[str(k)[:80]] = nums

    if grouped:
        # Collect all entity names that appear in the data
        seen_ents: list[str] = []
        for sub in grouped.values():
            for e in sub:
                if e not in seen_ents:
                    seen_ents.append(e)
        # Build colour map
        entity_colors = {}
        for i, e in enumerate(seen_ents):
            entity_colors[e] = _ENTITY_PALETTE[i % len(_ENTITY_PALETTE)]
        return _make_grouped_chart(grouped, title, units, entity_colors)

    # --- Flat mode: simple key→number -----------------------------------------
    clean = {}
    for k, v in chart_data.items():
        if isinstance(v, list):
            if all(isinstance(i, (int, float)) for i in v):
                v = v[-1]
            else:
                continue
        try:
            num = float(v)
            if not math.isnan(num) and not math.isinf(num):
                clean[str(k)[:80]] = num
        except Exception:
            continue

    if not clean:
        logger.warning("No valid numeric data to plot for chart: %s", title)
        return None

    labels = list(clean.keys())
    values = list(clean.values())

    max_label_length = max(len(label) for label in labels) if labels else 0
    if len(clean) > 12:
        horizontal = True
    elif max_label_length > 40:
        horizontal = True
    elif len(clean) <= 4 and max_label_length <= 20:
        horizontal = False
    elif len(clean) <= 6 and max_label_length <= 30:
        horizontal = False
    else:
        horizontal = True

    fig, ax = plt.subplots(figsize=(12, 8))

    if horizontal:
        wrapped_labels = ['\n'.join(textwrap.wrap(label, width=40)) for label in labels]
        colors = [_color_for_label(l, entities) for l in labels]
        bars = ax.barh(wrapped_labels, values, color=colors)
        ax.set_xlabel(units)
        ax.set_ylabel("Category")
        for bar in bars:
            width = bar.get_width()
            formatted_value = format_value_with_units(width, units)
            ax.annotate(formatted_value, xy=(width, bar.get_y() + bar.get_height() / 2),
                        xytext=(5, 0), textcoords="offset points",
                        ha='left', va='center', fontsize=8)
    else:
        wrapped_labels = ['\n'.join(textwrap.wrap(label, width=15)) for label in labels]
        colors = [_color_for_label(l, entities) for l in labels]
        bars = ax.bar(range(len(labels)), values, color=colors)
        ax.set_ylabel(units)
        ax.set_xlabel("Category")
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(wrapped_labels, ha='center', va='top')
        for bar in bars:
            height = bar.get_height()
            formatted_value = format_value_with_units(height, units)
            ax.annotate(formatted_value, xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 5), textcoords="offset points",
                        ha='center', va='bottom', fontsize=8)

    ax.set_title(title[:100])
    ax.grid(axis="y" if not horizontal else "x", linestyle="--", alpha=0.6)
    fig.tight_layout()

    filename = f"chart_{uuid.uuid4().hex}.png"
    path = os.path.join("charts", filename)
    fig.savefig(path, dpi=300, bbox_inches='tight')
    plt.close(fig)
    return path


def _chart_fingerprint(chart_data: dict) -> str:
    """Return a deterministic SHA-256 hex of *chart_data* for dedup purposes."""
    import hashlib, json

    def _flatten(obj):
        if isinstance(obj, dict):
            return {str(k): _flatten(v) for k, v in sorted(obj.items())}
        if isinstance(obj, (list, tuple)):
            return [_flatten(i) for i in obj]
        return obj

    canonical = json.dumps(_flatten(chart_data), sort_keys=True, default=str)
    return hashlib.sha256(canonical.encode()).hexdigest()


def append_to_doc(doc, section_data: dict, level: int = 2, citation_map: dict | None = None, skip_charts: bool = False, rendered_charts: set | None = None):
    """Append section to document with heading, paragraph, table, chart, and citations."""
    heading = section_data.get("heading", "Untitled Section")
    doc.add_heading(heading, level=level)

    text = section_data.get("text", "").strip()

    # Add citations to the text if sources are available
    if text and citation_map and section_data.get("sources"):
        citation_numbers = []
        for source in section_data.get("sources", []):
            source_key = f"{source.get('file', 'Unknown')}_{source.get('sheet', '')}_{source.get('entity', '')}"
            if source_key in citation_map:
                citation_numbers.append(citation_map[source_key])
        if citation_numbers:
            unique_citations = sorted(set(citation_numbers))
            citations_str = " " + "".join([f"[{num}]" for num in unique_citations])
            text = text + citations_str

    if text:
        add_inline_markdown_paragraph(doc, text)

    table_data = section_data.get("table", [])
    if isinstance(table_data, dict):
        table_data = [table_data]
    if isinstance(table_data, list) and table_data:
        add_table(doc, table_data)

    # Qualitative findings: policies, programmes, framework alignment — the material
    # the metrics table cannot represent.
    findings = section_data.get("findings", [])
    if isinstance(findings, list) and findings:
        for finding in findings:
            finding = str(finding).strip()
            if finding:
                add_inline_markdown_paragraph(doc, finding, style="List Bullet")

    # Skip charts if requested (for multi-vendor comparisons)
    if not skip_charts:
        chart_data = section_data.get("chart_data", {})
        if isinstance(chart_data, dict) and chart_data:
            # Deduplicate: skip if we already rendered an identical chart
            fp = _chart_fingerprint(chart_data)
            if rendered_charts is not None and fp in rendered_charts:
                logger.info(f"📊 Skipping duplicate chart for '{heading}'")
            else:
                if rendered_charts is not None:
                    rendered_charts.add(fp)

                entities = section_data.get("entities")
                units = section_data.get("units")
                chart_path = make_chart(chart_data, title=heading, entities=entities, units=units)
                if chart_path:
                    doc.add_picture(chart_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = 1  # center
                    logger.info(f"📊 Chart rendered for '{heading}'")
                    # Surface it in the UI too — charts were only ever visible
                    # inside the finished .docx.
                    progress_bus.publish_chart(chart_path)
                else:
                    logger.warning(f"📊 Chart generation returned None for '{heading}'")
        else:
            logger.info(f"📊 No chart_data for section '{heading}'")


def save_doc(doc, filename: str = "_report.docx"):
    """Save the Word document."""
    doc.save(filename)
    logger.info(f"✅ Report saved: {filename}")


class SectionWriterAgent:
    def __init__(self, llm, tokenizer=None):
        self.llm = llm
        self.tokenizer = tokenizer
        # Diagnostics, not narrative — these were print(), so they bypassed logging
        # entirely and appeared on screen regardless of log level.
        if tokenizer:
            logger.debug("Tokenizer initialized for SectionWriterAgent")
        else:
            logger.debug("No tokenizer provided for SectionWriterAgent")

    # Appended to the prompt on a retry. The failures seen in practice are malformed
    # JSON (unescaped quotes, possessives written as Supremo1"s), not the model being
    # unable to do the task — so restating the format constraint is usually enough.
    _RETRY_SUFFIX = """

IMPORTANT — your previous response could not be parsed as JSON. Return ONLY a single
valid JSON object. Escape any double quote inside a string value as \\". Do not use
apostrophes or possessive forms in any string value (write "Supremo1 targets" rather
than "Supremo1's targets"). Do not wrap the JSON in code fences or commentary."""

    @staticmethod
    def _default_attempts() -> int:
        """
        How many times to try parsing a section's JSON.

        Default 1 — no retry. A retry costs a full regeneration of a ~7k-token
        section prompt (~80s measured), which is too slow for a live demo: two
        retries in one run pushed a 136s report to over 4 minutes. Set
        SECTION_WRITE_ATTEMPTS=2 for unattended runs where completeness matters
        more than latency.
        """
        try:
            return max(1, int(os.environ.get("SECTION_WRITE_ATTEMPTS", "1")))
        except ValueError:
            return 1

    def _invoke_and_parse(
        self,
        prompt: str,
        label: str,
        entities: Optional[List[str]] = None,
        attempts: Optional[int] = None,
        expected_structure: Optional[str] = None,
    ) -> dict:
        """
        Call the LLM and parse its JSON, retrying once on a parse failure.

        Section writers previously had no retry: a single malformed response cost the
        entire section, which then rendered as an empty heading. One retry recovers
        nearly all of these, since the failure is formatting rather than capability.

        Passing `entities` matters — the possessive repair keys off it, and it was
        being called without them.
        """
        # Imported here, not at module scope: agent_factory imports this module, so a
        # top-level import would be circular. Every caller does the same.
        from agents.agent_factory import UniversalJSONCleaner

        if attempts is None:
            attempts = self._default_attempts()

        last_error: Optional[Exception] = None
        for attempt in range(1, attempts + 1):
            try:
                attempt_prompt = prompt if attempt == 1 else prompt + self._RETRY_SUFFIX
                response = self.llm.invoke([HumanMessage(content=attempt_prompt)]).content.strip()
                json_str = UniversalJSONCleaner.clean_and_extract_json(
                    response, expected_type="object", entities=entities
                )
                parsed = UniversalJSONCleaner.parse_with_validation(
                    json_str, expected_structure=expected_structure, entities=entities
                )
                if attempt > 1:
                    logger.info(f"✅ '{label}' recovered on attempt {attempt}")
                return parsed
            except Exception as e:
                last_error = e
                if attempt < attempts:
                    logger.warning(
                        f"⚠️ '{label}' attempt {attempt}/{attempts} did not parse ({e}); retrying"
                    )

        assert last_error is not None
        raise last_error

    def estimate_tokens(self, text: str) -> int:
        return max(1, len(text) // 4)

    def log_token_count(self, text: str, tokenizer=None, label: str = "Prompt"):
        if not text:
            logger.debug(f"Cannot log tokens: empty text for {label}")
            return
        if tokenizer:
            token_count = len(tokenizer.encode(text))
        else:
            token_count = self.estimate_tokens(text)
        logger.debug(f"{label} token count: {token_count}")

    def write_section(self, section_title: str, context_chunks: list[dict]) -> dict:
        from collections import defaultdict

        grouped = defaultdict(list)
        grouped_metadata = defaultdict(list)
        for chunk in context_chunks:
            entity = chunk.get("_search_entity", "Unknown")
            grouped[entity].append(chunk.get("content", ""))
            metadata = chunk.get("metadata", {})
            grouped_metadata[entity].append(metadata)

        entities = list(grouped.keys())
        
        # Handle different numbers of entities
        if len(entities) > 10:
            logger.warning(f"⚠️ Too many entities ({len(entities)}). Limiting to first 10.")
            entities = entities[:10]
        
        if len(entities) >= 2:
            # Multi-entity comparison (2-10 entities)
            return self._write_multi_entity_comparison(section_title, grouped, entities, grouped_metadata)
        elif len(entities) == 1:
            return self._write_single_entity_section(section_title, grouped, entities[0], grouped_metadata)
        else:
            logger.warning(f"⚠️ No valid entities found for section: {section_title}")
            return {
                "heading": section_title,
                "text": f"Insufficient data for analysis. Entities: {entities}",
                "table": [],
                "chart_data": {},
                "sources": [],
                # propagate for downstream report logic
                "is_comparison": False,
                "entities": entities
            }

    def write_section_typed(
        self,
        topic: str,
        chunks: List[Chunk],
        *,
        entities: Optional[List[str]] = None,
        is_comparison: bool = False,
    ) -> SectionDraft:
        """Write a section and return a typed SectionDraft model."""
        # Convert Chunk models to legacy dicts for the existing write_section
        legacy_chunks = [c.to_legacy_dict() for c in chunks]

        # Ensure _search_entity is set so write_section can group correctly.
        if entities and len(entities) == 1:
            for lc in legacy_chunks:
                lc.setdefault("_search_entity", entities[0])
        elif entities and len(entities) >= 2:
            # For multi-entity comparison: honour existing search_entity tags
            # first, then distribute untagged chunks round-robin across entities
            # so the comparison writer sees data for every entity.
            untagged = [lc for lc in legacy_chunks if not lc.get("_search_entity")]
            for i, lc in enumerate(untagged):
                lc["_search_entity"] = entities[i % len(entities)]

        result = self.write_section(topic, legacy_chunks)
        result.setdefault("entities", entities or [])
        result.setdefault("is_comparison", is_comparison)
        result["chunks_used"] = len(chunks)

        return SectionDraft.from_legacy_dict(result)

    def write_derived_section_typed(
        self,
        topic: str,
        role: str,
        prior_sections: List[SectionDraft],
        *,
        entities: Optional[List[str]] = None,
        query: str = "",
    ) -> SectionDraft:
        """
        Write a section that reasons over the already-written comparison sections
        instead of retrieving source data.

        role="synthesize" -> summary/assessment across sections
        role="recommend"  -> actions derived from the gaps those sections expose
        """
        entities = entities or []
        entity_label = " and ".join(entities) if entities else "the subjects"

        # Flatten what the compare sections actually established
        digest_parts = []
        for s in prior_sections:
            lines = [f"## {s.topic}"]
            if s.markdown:
                lines.append(s.markdown.strip())
            for row in s.table[:20]:
                metric = row.get("Metric", "")
                if not metric:
                    continue
                vals = " | ".join(f"{e}: {row.get(e, 'N/A')}" for e in entities) if entities else ""
                # Carry the Analysis verdict through. The compare section already
                # decided who leads; without it the synthesis LLM re-derives the
                # comparison from raw numbers and can invert it.
                verdict = str(row.get("Analysis", "")).strip()
                line = f"- {metric} — {vals}".rstrip(" —")
                # Deterministic direction, computed in code. This is what the
                # synthesis layer must cite instead of comparing numbers itself.
                higher = _higher_entity(row, list(entities))
                if higher:
                    line += f"  [LARGER VALUE: {higher}]"
                if verdict and verdict != "N/A":
                    line += f"  [section verdict: {verdict}]"
                lines.append(line)
            for f in s.findings[:10]:
                lines.append(f"- {f}")
            digest_parts.append("\n".join(lines))
        digest = "\n\n".join(digest_parts)

        if role == "recommend":
            objective = (
                f"Produce concrete, actionable recommendations for {entity_label}, derived from "
                "the gaps, shortfalls and weaker-performing areas visible in the SECTIONS below.\n"
                "Address each subject separately. Every recommendation must be traceable to a "
                "specific gap in the SECTIONS — a missed target, a slower pace, a weaker "
                "commitment, or an area where one subject trails the other.\n"
                "Do NOT restate metrics as if they were recommendations."
            )
            findings_instruction = (
                'findings: the recommendations themselves, one action per string, each naming '
                'the subject it applies to and the gap it addresses.'
            )
        else:
            objective = (
                f"Produce a summary assessment across all SECTIONS below for {entity_label}.\n"
                "State where each subject leads, trails, or matches, and whether it meets, "
                "exceeds or falls short of any target or external standard named in the SECTIONS.\n"
                "Do NOT introduce data that does not appear in the SECTIONS."
            )
            findings_instruction = (
                'findings: the key cross-cutting judgements, one per string, each naming the '
                'subject(s) it concerns.'
            )

        prompt = f"""You are writing the "{topic}" section of a report about {entity_label}.

OBJECTIVE:
{objective}

USER REQUEST (for tone and scope):
{query[:1500]}

SECTIONS ALREADY WRITTEN:
{digest}

Return JSON with exactly these keys:
- heading: a short descriptive title
- text: 3-5 sentences of narrative
- {findings_instruction}

RULES:
- Use only what appears in SECTIONS above; invent no figures.
- If the SECTIONS disagree about a value, say so rather than picking one silently.
- No possessive apostrophes (write "Oracle revenue", not "Oracle's revenue").
- Respond only in valid JSON.

DIRECTION OF COMPARISON — read carefully:
Do NOT work out for yourself which subject is ahead on a metric. That has already
been determined and is annotated on each line:
- "[LARGER VALUE: X]" means X holds the numerically larger figure on that line.
- "[section verdict: ...]" is the assessment written by the section author.
Any claim that one subject leads, trails, exceeds, lags or falls short on a metric
must agree with those annotations. A line with no annotation is undetermined —
report the two figures without asserting a direction.
Note a larger figure is not automatically better: for a reduction or a shortfall,
the smaller figure may be the stronger result. Read the metric name before judging.
"""

        from agents.agent_factory import UniversalJSONCleaner

        try:
            parsed = self._invoke_and_parse(
                prompt,
                label=topic,
                entities=list(entities) if entities else None,
                expected_structure="Object with 'heading', 'text', and 'findings' keys",
            )
            if not isinstance(parsed, dict):
                raise ValueError(f"expected object, got {type(parsed).__name__}")

            raw_findings = parsed.get("findings", [])
            if isinstance(raw_findings, str):
                raw_findings = [raw_findings]
            elif isinstance(raw_findings, dict):
                raw_findings = [raw_findings]
            findings = []
            if isinstance(raw_findings, list):
                for f in raw_findings:
                    # Models frequently return {"Supremo1": "do X"} instead of a
                    # plain string; render those as "Supremo1: do X" rather than
                    # letting str() emit a Python dict literal into the document.
                    if isinstance(f, dict):
                        named = f.get("finding") or f.get("text") or f.get("recommendation")
                        if named:
                            f = named
                        else:
                            f = "; ".join(
                                f"{k}: {v}" for k, v in f.items() if str(v).strip()
                            )
                    f = str(f).strip()
                    if f:
                        findings.append(f)

            logger.info(f"🧩 {role} section '{topic}': {len(findings)} findings from "
                        f"{len(prior_sections)} prior sections")

            return SectionDraft.from_legacy_dict({
                "heading": parsed.get("heading", topic),
                "text": parsed.get("text", ""),
                "table": [],
                "findings": findings,
                "chart_data": {},
                "entities": entities,
                # A synthesis/recommendation section is still part of a comparison
                # report. Hardcoding False here made the report header read
                # "Report: Supremo1" whenever one of these sorted first.
                "is_comparison": len(entities) >= 2,
                "chunks_used": 0,
            })

        except Exception as e:
            logger.error(f"⚠️ Failed to write {role} section '{topic}': {e}")
            return SectionDraft.from_legacy_dict({
                "heading": topic,
                "text": "",
                "entities": entities,
            })

    def _write_single_entity_section(self, section_title: str, grouped_chunks: dict, entity: str, grouped_metadata: dict | None = None) -> dict:
        text = "\n\n".join(grouped_chunks[entity])

        # Extract unique sources from metadata
        sources = []
        if grouped_metadata and entity in grouped_metadata:
            seen_sources = set()
            for metadata in grouped_metadata[entity]:
                source_key = f"{metadata.get('source', 'Unknown')}_{metadata.get('sheet', '')}"
                if source_key not in seen_sources:
                    sources.append({
                        "file": metadata.get("source", "Unknown"),
                        "sheet": metadata.get("sheet", ""),
                        "entity": entity
                    })
                    seen_sources.add(source_key)

        prompt = f"""Extract key data for {entity} on {section_title}.

Return JSON:
{{
  "heading": "descriptive title",
  "text": "2-sentence summary",
  "table": [{{"Metric": "metric name", "Value": "value with units"}}],
  "chart_data": {{"metric_name": numeric_value}}
}}

IMPORTANT: chart_data MUST contain only numeric values (no text, no units, no currency symbols).
Example chart_data: {{"Total Revenue": 14234, "Operating Income": 5831, "Net Income": 3926}}

Data:
{text[:2000]}

CRITICAL RULES:
1. NEVER use possessive forms or apostrophes (no 's).
   - Wrong: "Oracle's revenue", "company's growth"
   - Right: "Oracle revenue", "company growth", "revenue of Oracle"
2. Use "N/A" for missing data in tables.
3. Return valid JSON only - no apostrophes in text values.
4. Include at least 3-5 numeric metrics in chart_data if available."""

        try:
            self.log_token_count(prompt, self.tokenizer, label=f"SingleEntity Prompt ({section_title})")

            import ast

            parsed = self._invoke_and_parse(
                prompt,
                label=section_title,
                entities=[entity] if entity else None,
                expected_structure="Object with 'heading', 'text', 'table', and 'chart_data' keys",
            )

            chart_data = parsed.get("chart_data", {})
            if isinstance(chart_data, str):
                try:
                    import ast as _ast
                    chart_data = _ast.literal_eval(chart_data)
                except Exception:
                    chart_data = {}

            table = parsed.get("table", [])
            if isinstance(table, str):
                try:
                    import ast as _ast
                    table = _ast.literal_eval(table)
                except Exception:
                    table = []

            # Validate chart_data has actual numeric values
            has_numeric = any(
                _parse_numeric(v) is not None for v in chart_data.values()
            ) if isinstance(chart_data, dict) else False

            if not has_numeric and isinstance(table, list) and table:
                mined = _mine_chart_from_table(table)
                if mined:
                    logger.info(f"📊 Using table-mined chart data for '{section_title}' ({len(mined)} metrics)")
                    chart_data = mined

            return {
                "heading": parsed.get("heading", section_title),
                "text": parsed.get("text", ""),
                "table": table,
                "chart_data": chart_data,
                "sources": sources,
                "is_comparison": False,
                "entities": [entity]
            }

        except Exception as e:
            logger.error("❌ Failed to write single-entity section: %s", e)
            return {
                "heading": section_title,
                "text": f"Could not generate section due to error: {e}",
                "table": [],
                "chart_data": {},
                "sources": sources,
                "is_comparison": False,
                "entities": [entity]
            }

    def _write_multi_entity_comparison(self, section_title: str, grouped_chunks: dict, entities: list[str], grouped_metadata: dict | None = None) -> dict:
        """Handle comparisons between 2-10 entities with special formatting for tender responses"""
        import ast
        from agents.agent_factory import UniversalJSONCleaner
        
        # For backward compatibility, use the old method for exactly 2 entities
        if len(entities) == 2:
            return self._write_comparison_section(section_title, grouped_chunks, entities, grouped_metadata)
        
        # Multi-entity comparison (3-10 entities)
        entity_data_sections = []
        for entity in entities:
            entity_text = "\n\n".join(grouped_chunks.get(entity, []))[:1500]  # Limit text per entity
            entity_data_sections.append(f"=== {entity} ===\n{entity_text}")
        
        entities_str = ", ".join(entities)
        data_section = "\n\n".join(entity_data_sections)
        
        # Detect if this is a tender/RFP comparison
        is_tender = "tender" in section_title.lower() or "rfp" in section_title.lower() or "proposal" in section_title.lower()
        
        if is_tender:
            table_instruction = f"""
- table: List of evaluation criteria as rows, with columns for each entity
  Format: {{"Metric": "criterion", "{entities[0]}": "value1", "{entities[1]}": "value2", ...}}
- Include a "Best Value" or "Ranking" column if comparing quantitative metrics
- For qualitative criteria, use ratings like "Excellent", "Good", "Fair", "Poor"
"""
        else:
            table_instruction = f"""
- table: List of metrics as rows, with columns for each entity
  Format: {{"Metric": "metric_name", "{entities[0]}": "value1", "{entities[1]}": "value2", ...}}
- Include an "Analysis" column highlighting key differences
"""

        prompt = f"""
You are writing a structured comparison section for {len(entities)} entities: {entities_str}.

Topic: {section_title}

OBJECTIVE:
Create a comprehensive comparison table showing all entities side-by-side.

Always follow this exact structure in your JSON output:
- heading: A descriptive title for the section
- text: A 2-3 sentence overview comparing all entities
{table_instruction}
- chart_data: Comparable numeric values from all entities

DATA:
{data_section}

INSTRUCTIONS:
- Extract specific metrics that can be compared across all entities
- Use "N/A" if an entity is missing a value
- For tender/RFP comparisons, focus on evaluation criteria
- Keep values human-readable
- Ensure fair and balanced comparison

CRITICAL RULES:
1. NEVER use possessive forms (no 's)
2. Ensure valid JSON format
3. Include ALL entities in each table row

Respond only in valid JSON format.
"""

        try:
            parsed = self._invoke_and_parse(
                prompt,
                label=section_title,
                entities=list(entities) if entities else None,
                expected_structure="Object with 'heading', 'text', 'table', and 'chart_data' keys",
            )

            # Process table to ensure all entities have columns
            table = parsed.get("table", [])
            if isinstance(table, str):
                try:
                    table = ast.literal_eval(table)
                except:
                    table = []
            
            validated_table = []
            for row in table:
                if isinstance(row, dict):
                    validated_row = {"Metric": row.get("Metric", "Unknown")}
                    has_data = False
                    for entity in entities:
                        value = row.get(entity, "N/A")
                        validated_row[entity] = value
                        if value != "N/A":
                            has_data = True
                    # Add analysis/ranking columns if present
                    for key in ["Analysis", "Best Value", "Ranking"]:
                        if key in row:
                            validated_row[key] = row[key]
                    if has_data:
                        validated_table.append(validated_row)
            
            # Build chart_data with entity names in the keys so bars get
            # per-entity colors.  Prefer mining the validated table.
            chart_data = _chart_data_from_table(validated_table, entities)
            if not chart_data:
                raw_cd = parsed.get("chart_data", {})
                if isinstance(raw_cd, str):
                    try:
                        raw_cd = ast.literal_eval(raw_cd)
                    except Exception:
                        raw_cd = {}
                if isinstance(raw_cd, dict) and raw_cd:
                    chart_data = raw_cd
            if not chart_data:
                mined = _mine_chart_from_table(validated_table)
                if mined:
                    chart_data = mined
                    logger.info(f"📊 Using table-mined chart data for multi-entity '{section_title}'")

            # Extract sources
            sources = []
            if grouped_metadata:
                seen_sources = set()
                for entity in entities:
                    if entity in grouped_metadata:
                        for metadata in grouped_metadata[entity]:
                            source_key = f"{metadata.get('source', 'Unknown')}_{metadata.get('sheet', '')}_{entity}"
                            if source_key not in seen_sources:
                                sources.append({
                                    "file": metadata.get("source", "Unknown"),
                                    "sheet": metadata.get("sheet", ""),
                                    "entity": entity
                                })
                                seen_sources.add(source_key)
            
            return {
                "heading": parsed.get("heading", section_title),
                "text": parsed.get("text", ""),
                "table": validated_table,
                "chart_data": chart_data,
                "sources": sources,
                "is_comparison": True,
                "entities": entities,
                "is_tender": is_tender
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to write multi-entity comparison: {e}")
            return {
                "heading": section_title,
                "text": f"Could not generate comparison for {len(entities)} entities",
                "table": [],
                "chart_data": {},
                "sources": [],
                "is_comparison": True,
                "entities": entities
            }
    
    def _write_comparison_section(self, section_title: str, grouped_chunks: dict, entities: list[str], grouped_metadata: dict | None = None) -> dict:
        """Legacy method for exactly 2 entities - kept for backward compatibility"""
        import ast
        from agents.agent_factory import UniversalJSONCleaner

        if len(entities) != 2:
            raise ValueError("This method requires exactly two entities. Use _write_multi_entity_comparison for more.")

        entity_a, entity_b = entities
        text_a = "\n\n".join(grouped_chunks[entity_a])
        text_b = "\n\n".join(grouped_chunks[entity_b])

        prompt = f"""
You are writing a structured section for a comparison report between {entity_a} and {entity_b}.

Topic: {section_title}

OBJECTIVE:
Summarize key data from the context and produce a clear, side-by-side comparison table.

Always follow this exact structure in your JSON output:
- heading: A short, descriptive title for the section
- text: A 2–4 sentence overview comparing {entity_a} and {entity_b}
- table: List of dicts formatted as: Metric | {entity_a} | {entity_b} | Analysis
- findings: List of strings capturing material that is NOT a numeric metric
- chart_data: A dictionary of comparable numeric values to plot

DATA:
=== {entity_a} ===
{text_a}

=== {entity_b} ===
{text_b}

INSTRUCTIONS:
- Extract specific metrics (numbers, %, dates) from the data
- Use "N/A" if one entity is missing a value
- Use analysis terms like: "Higher", "Lower", "Similar", "{entity_a} only", "{entity_b} only"
- Do not echo file names or metadata
- Keep values human-readable (e.g., "18,500 tonnes CO2e")
- "N/A" means the value is absent from the DATA above. It does NOT mean the entity
  has no such commitment — never state or imply that an entity lacks a policy when
  its value is simply missing from the context.

FINDINGS (do not skip):
Much of the material above is qualitative and has no place in the metrics table.
Put it in "findings" instead of discarding it. Cover, where the DATA supports it:
- Policy commitments and restrictive financing rules (e.g. thermal coal, oil and gas
  by region or type, sector-level financed-emissions targets and their named sectors)
- Named programmes and initiatives
- Alignment with external frameworks and standards, and whether each entity meets,
  exceeds or falls short of a stated target or threshold
Each finding is one self-contained sentence naming the entity or entities it concerns.
Return [] only if the DATA genuinely contains no such material.

CRITICAL RULES:
1. NEVER use possessive forms or apostrophes (no 's).
   - Wrong: "Oracle's revenue", "company's performance"  
   - Right: "Oracle revenue", "company performance", "revenue of Oracle"
2. Ensure all JSON is valid - no apostrophes in text values.
3. Use proper escaping if quotes are needed in text.

Respond only in valid JSON format.
"""

        try:
            if self.tokenizer:
                self.log_token_count(prompt, self.tokenizer, label=f"Comparison Prompt ({section_title})")
            else:
                logger.warning("⚠️ No tokenizer available for token counting in SectionWriterAgent")

            parsed = self._invoke_and_parse(
                prompt,
                label=section_title,
                entities=entities,
                expected_structure="Object with 'heading', 'text', 'table', and 'chart_data' keys",
            )

            chart_data = parsed.get("chart_data", {})
            if isinstance(chart_data, str):
                try:
                    chart_data = ast.literal_eval(chart_data)
                except Exception as e:
                    logger.warning("⚠️ Failed to parse chart_data: %s", e)
                    chart_data = {}
            if not isinstance(chart_data, dict):
                chart_data = {}

            table = parsed.get("table", [])
            if isinstance(table, str):
                try:
                    table = ast.literal_eval(table)
                except Exception:
                    table = []
            if isinstance(table, dict):
                table = [table]
            elif not isinstance(table, list):
                table = []

            validated = []
            for row in table:
                if not isinstance(row, dict):
                    continue
                validated_row = {
                    "Metric": row.get("Metric", "Unknown Metric"),
                    entity_a: row.get(entity_a, "N/A"),
                    entity_b: row.get(entity_b, "N/A"),
                    "Analysis": row.get("Analysis", "N/A")
                }
                if validated_row[entity_a] != "N/A" or validated_row[entity_b] != "N/A":
                    validated.append(validated_row)

            # Build chart_data with entity names in the keys so bars get
            # per-entity colors.  Prefer mining the validated table (reliable)
            # over the LLM's freeform chart_data (often single-entity / flat).
            final_chart_data = _chart_data_from_table(validated, [entity_a, entity_b])
            if not final_chart_data:
                # Fallback 1: try the LLM's chart_data (may be nested or flat)
                if isinstance(chart_data, dict) and chart_data:
                    final_chart_data = chart_data
                    logger.info(f"📊 Using LLM chart_data fallback for '{section_title}'")
            if not final_chart_data:
                # Fallback 2: mine any numeric data from the table
                mined = _mine_chart_from_table(validated)
                if mined:
                    final_chart_data = mined
                    logger.info(f"📊 Using table-mined chart data for '{section_title}' ({len(mined)} metrics)")

            # Extract unique sources
            sources = []
            if grouped_metadata:
                seen_sources = set()
                for entity in entities:
                    if entity in grouped_metadata:
                        for metadata in grouped_metadata[entity]:
                            source_key = f"{metadata.get('source', 'Unknown')}_{metadata.get('sheet', '')}_{entity}"
                            if source_key not in seen_sources:
                                sources.append({
                                    "file": metadata.get("source", "Unknown"),
                                    "sheet": metadata.get("sheet", ""),
                                    "entity": entity
                                })
                                seen_sources.add(source_key)

            # Qualitative material that has no row in the metrics table
            raw_findings = parsed.get("findings", [])
            if isinstance(raw_findings, str):
                raw_findings = [raw_findings]
            findings = []
            if isinstance(raw_findings, list):
                for f in raw_findings:
                    if isinstance(f, dict):
                        f = f.get("finding") or f.get("text") or ""
                    f = str(f or "").strip()
                    if f:
                        findings.append(f)
            if findings:
                logger.info(f"📝 {len(findings)} qualitative findings for '{section_title}'")

            return {
                "heading": parsed.get("heading", section_title),
                "text": parsed.get("text", ""),
                "table": validated,
                "findings": findings,
                "chart_data": final_chart_data,
                "sources": sources,
                "is_comparison": True,
                "entities": [entity_a, entity_b]
            }

        except Exception as e:
            logger.error("⚠️ Failed to write comparison section: %s", e)
            sources = []
            if grouped_metadata:
                seen_sources = set()
                for entity in entities:
                    if entity in grouped_metadata:
                        for metadata in grouped_metadata[entity]:
                            source_key = f"{metadata.get('source', 'Unknown')}_{metadata.get('sheet', '')}_{entity}"
                            if source_key not in seen_sources:
                                sources.append({
                                    "file": metadata.get("source", "Unknown"),
                                    "sheet": metadata.get("sheet", ""),
                                    "entity": entity
                                })
                                seen_sources.add(source_key)

            return {
                "heading": section_title,
                "text": f"Could not generate summary due to error: {e}",
                "table": [],
                "chart_data": {},
                "sources": sources,
                "is_comparison": True,
                "entities": entities
            }


class ReportWriterAgent:
    def __init__(self, doc=None, model_name: str = "unknown", llm=None):
        # Derive model name from the LLM object when not explicitly provided
        if model_name == "unknown" and llm is not None:
            model_name = (
                getattr(llm, "model_name", None)
                or getattr(llm, "model_id", None)
                or getattr(llm, "model", None)
                or "unknown"
            )
        self.model_name = str(model_name)
        self.llm = llm  # Store LLM for generating summaries

    @staticmethod
    def _has_section_titled(sections: list[dict], *phrases: str) -> bool:
        """
        True if a planned section already covers one of these headings.

        Matched on exact heading or a "<phrase>:" / "<phrase> " prefix, so
        "Executive Summary: ESG Comparison of X and Y" counts but a section
        merely mentioning the words in passing does not.
        """
        for s in sections:
            if s.get("is_category_header"):
                continue
            heading = str(s.get("heading", "")).strip().lower()
            if not heading:
                continue
            for phrase in phrases:
                p = phrase.strip().lower()
                if heading == p or heading.startswith(p + ":") or heading.startswith(p + " "):
                    logger.info(f"Planned section '{s.get('heading')}' covers '{phrase}' — "
                                f"skipping the generated one")
                    return True
        return False

    def _generate_executive_summary(self, sections: list[dict], is_comparison: bool, entities: list[str], query: str | None = None) -> str:
        if not self.llm:
            return self._generate_intro_section(is_comparison, entities)

        section_summaries = []
        for section in sections:
            heading = section.get("heading", "Unknown Section")
            text = section.get("text", "")
            if text:
                section_summaries.append(f"{heading}: {text}")

        sections_text = "\n\n".join(section_summaries)

        query_context = f"\nUser's Original Request:\n{query}\n" if query else ""

        if is_comparison:
            # Handle 2-10 entities comparison
            if len(entities) == 2:
                entity_description = f"{entities[0]} and {entities[1]}"
            else:
                entity_description = ", ".join(entities[:-1]) + f", and {entities[-1]}"

            prompt = f"""
You are writing an executive summary for a comparison report between {len(entities)} entities: {entity_description}.
{query_context}
Based on the user's request and the following section summaries, create a 2-3 paragraph executive summary that:
1. Directly addresses what the user asked for
2. Highlights the most significant findings and differences across all {len(entities)} entities
3. Provides a clear overview of how the report answers their specific questions
4. For tender/RFP comparisons, identify the strongest candidates based on the evaluation criteria

Section Summaries:
{sections_text}

CRITICAL: Never use possessive forms (no apostrophes). Write "Oracle revenue" not "Oracle's revenue", "company performance" not "company's performance".

Write in a professional, analytical tone. Focus on answering the user's specific request.
"""
        else:
            prompt = f"""
You are writing an executive summary for a report about {entities[0] if entities else 'the organization'}.
{query_context}
Based on the user's request and the following section summaries, create a 2-3 paragraph executive summary that:
1. Directly addresses what the user asked for
2. Highlights the most significant findings relevant to their query
3. Provides a clear overview of how the report answers their specific questions

Section Summaries:
{sections_text}

CRITICAL: Never use possessive forms (no apostrophes). Write "Oracle revenue" not "Oracle's revenue", "company performance" not "company's performance".

Write in a professional, analytical tone. Focus on answering the user's specific request.
"""

        try:
            response = self.llm.invoke([HumanMessage(content=prompt)]).content.strip()
            return response
        except Exception as e:
            logger.warning(f"Failed to generate executive summary: {e}")
            return self._generate_intro_section(is_comparison, entities)

    def _generate_conclusion(self, sections: list[dict], is_comparison: bool, entities: list[str], query: str | None = None) -> str:
        if not self.llm:
            return "This analysis provides insights based on available data from retrieved documents."

        key_findings = []
        for section in sections:
            heading = section.get("heading", "Unknown Section")
            text = section.get("text", "")
            table = section.get("table", [])

            if table and isinstance(table, list):
                for row in table[:3]:
                    if isinstance(row, dict):
                        metric = row.get("Metric", "")
                        if metric:
                            key_findings.append(f"{heading}: {metric}")

            if text:
                key_findings.append(f"{heading}: {text}")

        findings_text = "\n".join(key_findings[:8])

        query_context = f"\nUser's Original Request:\n{query}\n" if query else ""

        if is_comparison:
            # Handle 2-10 entities comparison
            if len(entities) == 2:
                entity_description = f"{entities[0]} and {entities[1]}"
            else:
                entity_description = ", ".join(entities[:-1]) + f", and {entities[-1]}"

            prompt = f"""
Based on the analysis of {len(entities)} entities ({entity_description}), write a conclusion that directly answers the user's request.
{query_context}
Key Findings:
{findings_text}

Write 2-3 paragraphs that:
- Directly answer what the user asked for
- Summarize the main differences and similarities across all {len(entities)} entities
- Provide actionable insights based on their specific needs
- For tender/RFP comparisons, provide clear recommendations on which vendors best meet the requirements
- Include specific recommendations if appropriate

CRITICAL: Never use possessive forms (no apostrophes). Write "Oracle revenue" not "Oracle's revenue", "company growth" not "company's growth".

Focus on providing value for the user's specific use case.
"""
        else:
            prompt = f"""
Based on the analysis of {entities[0] if entities else 'the organization'}, write a conclusion that directly answers the user's request.
{query_context}
Key Findings:
{findings_text}

Write 2-3 paragraphs that:
- Directly answer what the user asked for
- Summarize the main insights relevant to their query
- Provide actionable insights based on their specific needs
- Include specific recommendations if appropriate

CRITICAL: Never use possessive forms (no apostrophes). Write "Oracle revenue" not "Oracle's revenue", "company growth" not "company's growth".

Focus on providing value for the user's specific use case.
"""

        try:
            response = self.llm.invoke([HumanMessage(content=prompt)]).content.strip()
            return response
        except Exception as e:
            logger.warning(f"Failed to generate conclusion: {e}")
            return "This analysis provides insights based on available data from retrieved documents."

    def _filter_failed_sections(self, sections: list[dict]) -> list[dict]:
        filtered_sections = []
        error_patterns = [
            "Could not generate",
            "due to error:",
            "Expecting ',' delimiter:",
            "Failed to",
            "Error:",
            "Exception:",
            "Traceback"
        ]
        for section in sections:
            text = section.get("text", "")
            heading = section.get("heading", "")
            has_error = any(pattern in text for pattern in error_patterns)
            if not has_error:
                filtered_sections.append(section)
            else:
                logger.info(f"🚫 Filtered out failed section: {heading}")
        return filtered_sections

    def _apply_document_styling(self, doc):
        from docx.shared import Pt, RGBColor

        _ORACLE_RED = RGBColor(0xC7, 0x46, 0x34)
        _CHARCOAL = RGBColor(0x31, 0x2D, 0x2A)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = _CHARCOAL

        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = _ORACLE_RED

        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = _CHARCOAL

    def _generate_report_title(self, is_comparison: bool, entities: list[str], query: str | None, sections: list[dict]) -> str:
        if query and self.llm:
            try:
                entity_context = f"{entities[0]} vs {entities[1]}" if is_comparison and len(entities) >= 2 else entities[0] if entities else "Organization"
                prompt = f"""Generate a concise, professional report title (max 10 words) based on:
User Query: {query}
Entities: {entity_context}
Type: {'Comparison' if is_comparison else 'Analysis'} Report

CRITICAL: Never use possessive forms (no apostrophes). Write "Oracle Performance" not "Oracle's Performance".

Return ONLY the title, no quotes or extra text."""
                title = self.llm.invoke([HumanMessage(content=prompt)]).content.strip()
                title = title.replace('"', '').replace("'", '').strip()
                if len(title) > 100:
                    title = title[:97] + "..."
                return title
            except Exception as e:
                logger.warning(f"Failed to generate dynamic title: {e}")

        if query:
            query_lower = query.lower()
            if "esg" in query_lower or "sustainability" in query_lower:
                topic_type = "ESG & Sustainability"
            elif "financial" in query_lower or "performance" in query_lower:
                topic_type = "Financial Performance"
            elif "risk" in query_lower:
                topic_type = "Risk Assessment"
            elif "governance" in query_lower:
                topic_type = "Corporate Governance"
            elif "climate" in query_lower or "carbon" in query_lower:
                topic_type = "Climate & Environmental"
            else:
                topic_type = "Business Analysis"
        else:
            section_topics = [s.get("heading", "") for s in sections[:3]]
            if any("climate" in h.lower() or "carbon" in h.lower() for h in section_topics):
                topic_type = "Climate & Environmental"
            elif any("esg" in h.lower() or "sustainability" in h.lower() for h in section_topics):
                topic_type = "ESG & Sustainability"
            else:
                topic_type = "Business Analysis"

        if is_comparison and len(entities) >= 2:
            return f"{topic_type} Report: {entities[0]} vs {entities[1]}"
        elif entities:
            return f"{topic_type} Report: {entities[0]}"
        else:
            return f"{topic_type} Report"

    def _add_report_header(self, doc, report_title: str, is_comparison: bool, entities: list[str]):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        _ORACLE_RED = RGBColor(0xC7, 0x46, 0x34)
        _CHARCOAL = RGBColor(0x31, 0x2D, 0x2A)
        _MID_GREY = RGBColor(0x74, 0x74, 0x74)

        title_paragraph = doc.add_heading(report_title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if is_comparison and len(entities) >= 2:
            if len(entities) == 2:
                subtitle = f"Comparative Analysis: {entities[0]} and {entities[1]}"
            else:
                subtitle = f"Comparative Analysis: {', '.join(entities[:-1])} and {entities[-1]}"
        elif entities:
            subtitle = f"Analysis of {entities[0]}"
        else:
            subtitle = "Comprehensive Analysis Report"

        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = "Calibri"
        subtitle_run.font.color.rgb = _CHARCOAL
        subtitle_run.italic = True

        now = datetime.datetime.now()
        date_str = now.strftime("%B %d, %Y")
        time_str = now.strftime("%H:%M")

        doc.add_paragraph()
        metadata_paragraph = doc.add_paragraph()
        metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_text = f"Generated on {date_str} at {time_str}\nModel: {self.model_name} | Powered by OCI Generative AI"
        metadata_run = metadata_paragraph.add_run(metadata_text)
        metadata_run.font.size = Pt(9)
        metadata_run.font.name = "Calibri"
        metadata_run.font.color.rgb = _MID_GREY

        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run("─" * 50)
        sep_run.font.color.rgb = _ORACLE_RED
        sep_run.font.size = Pt(8)
        doc.add_paragraph()

    def _generate_intro_section(self, is_comparison: bool, entities: list[str]) -> str:
        if is_comparison:
            comparison_note = f"This report compares data between {entities[0]} and {entities[1]} across key topics."
        else:
            comparison_note = f"This report presents information for {entities[0]}."
        return (
            f"{comparison_note} All data is sourced from retrieved documents and structured using LLM-based analysis.\n\n"
            "The analysis includes tables and charts where possible. Missing data is noted explicitly."
        )

    def _organize_sections(self, sections: list[dict], query: str | None, entities: list[str]) -> list[dict]:
        """
        Order sections for the document.

        The LLM-driven ordering is off by default. It has been failing on every
        observed run with `unsupported operand type(s) for -: 'dict' and 'int'`
        (_organize_sections_with_llm assumes the model returns section indices, but it
        returns objects), then silently falling back to the flat order used here — so
        it was costing a full serial round trip and an error line on screen to produce
        the result we get for free. Set REPORT_ORGANIZE_SECTIONS=1 to re-enable, and
        see §G of docs/rag_pipeline_fixes_2026-07-19.md for the underlying bug.
        """
        if os.environ.get("REPORT_ORGANIZE_SECTIONS", "").lower() in ("1", "true", "yes"):
            return self._organize_sections_with_llm(sections, query, entities)

        logger.debug("Section organisation: using planner order (LLM ordering disabled)")
        return sections

    def _organize_sections_with_llm(self, sections: list[dict], query: str | None, entities: list[str]) -> list[dict]:
        if not query or not self.llm or not sections:
            return sections
        section_info = []
        for i, section in enumerate(sections):
            section_info.append(f"{i+1}. {section.get('heading', 'Untitled Section')}")
        sections_list = "\n".join(section_info)

        prompt = f"""You are organizing sections for a report about {', '.join(entities)}.

User's Original Request:
{query}

Available Sections (numbered):
{sections_list}

Based on the user's request, create a hierarchical structure for these sections. The user's request may contain numbered main categories (like 1) Climate Impact, 2) Social Impact, etc.).

Return a JSON structure that organizes these sections hierarchically. Use the section numbers to reference them.

Format:
{{
  "structure": [
    {{
      "title": "Main Category Title from User's Request",
      "level": 1,
      "sections": [1, 3, 5]
    }},
    {{
      "title": "Another Main Category",
      "level": 1,
      "sections": [2, 4, 6]
    }}
  ],
  "orphan_sections": [7, 8]
}}

IMPORTANT:
- Extract main category titles from the user's request if they provided structured sections
- Group related sections under appropriate main categories
- Use level 1 for main categories, sections will be level 2
- List any sections that don't fit as orphan_sections
- Use the exact section numbers from the list above

Return ONLY valid JSON."""

        try:
            response = self.llm.invoke([HumanMessage(content=prompt)]).content.strip()
            import json, re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                structure = json.loads(json_str)

                organized = []
                used_sections = set()

                for category in structure.get("structure", []):
                    organized.append({
                        "heading": category.get("title", "Category"),
                        "level": 1,
                        "is_category_header": True
                    })
                    for section_num in category.get("sections", []):
                        idx = section_num - 1
                        if 0 <= idx < len(sections) and idx not in used_sections:
                            section_copy = sections[idx].copy()
                            section_copy["level"] = 2
                            organized.append(section_copy)
                            used_sections.add(idx)

                for section_num in structure.get("orphan_sections", []):
                    idx = section_num - 1
                    if 0 <= idx < len(sections) and idx not in used_sections:
                        section_copy = sections[idx].copy()
                        section_copy["level"] = 2
                        organized.append(section_copy)
                        used_sections.add(idx)

                for i, section in enumerate(sections):
                    if i not in used_sections:
                        section_copy = section.copy()
                        section_copy["level"] = 2
                        organized.append(section_copy)

                return organized
        except Exception as e:
            logger.warning(f"Failed to organize sections with LLM: {e}")

        return sections

    def _build_references_section(self, sections: list[dict]) -> tuple[dict, str]:
        all_sources = []
        citation_map = {}
        citation_counter = 1
        seen_sources = set()
        for section in sections:
            sources = section.get("sources", [])
            for source in sources:
                source_key = f"{source.get('file', 'Unknown')}_{source.get('sheet', '')}_{source.get('entity', '')}"
                if source_key not in seen_sources:
                    all_sources.append(source)
                    citation_map[source_key] = citation_counter
                    citation_counter += 1
                    seen_sources.add(source_key)

        references_text = []
        for i, source in enumerate(all_sources, 1):
            file_name = source.get("file", "Unknown")
            sheet = source.get("sheet", "")
            entity = source.get("entity", "")
            if sheet:
                ref_text = f"[{i}] {file_name}, Sheet: {sheet}"
            else:
                ref_text = f"[{i}] {file_name}"
            if entity:
                ref_text += f" ({entity})"
            references_text.append(ref_text)

        return citation_map, "\n".join(references_text)
    
    def _is_multi_vendor_comparison(self, sections: list[dict]) -> bool:
        """Check if this is a multi-vendor comparison (n > 2 vendors)"""
        # Check the first few sections to determine if this is a vendor comparison
        for section in sections[:3]:  # Check first 3 sections
            entities = section.get("entities", [])
            is_comparison = section.get("is_comparison", False)
            is_tender = section.get("is_tender", False)
            
            # Multi-vendor comparison if:
            # 1. More than 2 entities AND
            # 2. It's a comparison AND
            # 3. It's a tender/RFP/vendor comparison
            if len(entities) > 2 and is_comparison and is_tender:
                return True
                
            # Also check section headings for vendor/tender/RFP keywords
            heading = section.get("heading", "").lower()
            if len(entities) > 2 and is_comparison:
                vendor_keywords = ["vendor", "tender", "rfp", "proposal", "bid", "supplier", "quotation"]
                if any(keyword in heading for keyword in vendor_keywords):
                    return True
        
        return False
    




    def _create_vendor_comparison_visualization(self, sections: list[dict]) -> str | None:
        """Create a comprehensive visualization table for multi-vendor comparisons"""
        try:
            # Extract vendor names and metrics from all sections
            vendors = set()
            all_metrics = {}
            
            for section in sections:
                entities = section.get("entities", [])
                if entities and len(entities) > 2:
                    vendors.update(entities)
                
                # Extract metrics from tables
                table_data = section.get("table", [])
                if isinstance(table_data, list):
                    for row in table_data:
                        if isinstance(row, dict) and "Metric" in row:
                            metric_name = row["Metric"]
                            if metric_name not in all_metrics:
                                all_metrics[metric_name] = {}
                            
                            for vendor in entities:
                                if vendor in row:
                                    all_metrics[metric_name][vendor] = row[vendor]
            
            if not vendors or not all_metrics:
                logger.warning("No vendor data found for visualization")
                return None
            
            vendors = sorted(list(vendors))
            
            # Create the visualization
            fig, ax = plt.subplots(figsize=(14, max(8, len(all_metrics) * 0.5 + 2)))
            ax.axis('tight')
            ax.axis('off')
            
            # Prepare table data
            table_headers = ["Evaluation Criteria"] + vendors
            table_rows = []
            
            # Color mapping for ratings
            color_map = {
                'green': '#90EE90',      # Light green
                'yellow': '#FFFFE0',     # Light yellow
                'red': '#FFB6C1',        # Light red
                'excellent': '#90EE90',
                'good': '#B4EEB4',
                'fair': '#FFFFE0',
                'poor': '#FFB6C1',
                'best': '#90EE90',
                'n/a': '#F0F0F0'
            }

            # Determine wrap width dynamically
            wrap_width = int(25 * (14 / fig.get_size_inches()[0]))
            
            # Process metrics and create rows
            for metric, vendor_values in all_metrics.items():
                wrapped_metric = textwrap.fill(metric, wrap_width)
                row = [wrapped_metric]
                row_colors = ['#E6E6FA']  # Lavender for metric column
                
                for vendor in vendors:
                    value = vendor_values.get(vendor, "N/A")
                    wrapped_value = textwrap.fill(str(value), wrap_width)
                    row.append(wrapped_value)
                    
                    # Determine cell color based on value
                    value_lower = str(value).lower()
                    cell_color = '#FFFFFF'  # Default white
                    
                    # Check for color-coded words
                    for keyword, color in color_map.items():
                        if keyword in value_lower:
                            cell_color = color
                            break
                    
                    # Check for numeric comparisons
                    if cell_color == '#FFFFFF' and value != "N/A":
                        try:
                            numeric_value = float(str(value).replace('%', '').replace(',', ''))
                            all_nums = []
                            for v in vendor_values.values():
                                try:
                                    num = float(str(v).replace('%', '').replace(',', ''))
                                    all_nums.append(num)
                                except:
                                    pass
                            if all_nums:
                                min_val = min(all_nums)
                                max_val = max(all_nums)
                                if max_val > min_val:
                                    norm_value = (numeric_value - min_val) / (max_val - min_val)
                                    if norm_value > 0.66:
                                        cell_color = '#90EE90'
                                    elif norm_value > 0.33:
                                        cell_color = '#FFFFE0'
                                    else:
                                        cell_color = '#FFB6C1'
                        except:
                            pass
                    
                    row_colors.append(cell_color)
                
                table_rows.append((row, row_colors))
            
            # Create the table
            table_data = []
            cell_colors = []
            
            # Add header row
            header_colors = ['#4472C4'] * len(table_headers)
            table_data.append([textwrap.fill(h, wrap_width) for h in table_headers])
            cell_colors.append(header_colors)
            
            # Add data rows
            for row, colors in table_rows:
                table_data.append(row)
                cell_colors.append(colors)
            
            # Create table with styling
            table = ax.table(
                cellText=table_data,
                cellLoc='center',
                loc='center',
                colWidths=[0.25] + [0.15] * len(vendors),
                cellColours=cell_colors
            )
            
            # Style the table
            table.auto_set_font_size(False)
            table.set_fontsize(9)
            table.scale(1, 2.2)  # slightly taller for wrapped text
            
            # Bold headers
            for i in range(len(table_headers)):
                cell = table[(0, i)]
                cell.set_text_props(weight='bold', color='white')
                cell.set_height(0.1)
            
            # Set row heights and text properties
            for i in range(1, len(table_data)):
                for j in range(len(table_headers)):
                    cell = table[(i, j)]
                    cell.set_height(0.08)
                    if j == 0:  # Metric column
                        cell.set_text_props(weight='bold')
        
            
            # Add legend
            legend_elements = [
                mpatches.Patch(color='#90EE90', label='Excellent/Best'),
                mpatches.Patch(color='#B4EEB4', label='Good'),
                mpatches.Patch(color='#FFFFE0', label='Fair/Average'),
                mpatches.Patch(color='#FFB6C1', label='Poor/Below Average'),
                mpatches.Patch(color='#F0F0F0', label='N/A')
            ]
            ax.legend(handles=legend_elements, loc='upper center', bbox_to_anchor=(0.5, -0.05),
                    ncol=5, frameon=False, fontsize=8)
            
            
            # Save the figure
            filename = f"vendor_comparison_matrix_{uuid.uuid4().hex}.png"
            os.makedirs("charts", exist_ok=True)
            path = os.path.join("charts", filename)
            fig.savefig(path, dpi=300, bbox_inches='tight', pad_inches=0.5)
            plt.close(fig)
            
            logger.info(f"✅ Created vendor comparison visualization: {path}")
            return path
            
        except Exception as e:
            logger.error(f"❌ Failed to create vendor comparison visualization: {e}")
            return None


    def write_report(self, sections: list[dict], filter_failures: bool = True, query: str | None = None, output_dir: str = "reports") -> str:
        if not isinstance(sections, list):
            raise TypeError("Expected list of sections")

        if filter_failures:
            sections = self._filter_failed_sections(sections)
            logger.info(f"📊 After filtering failures: {len(sections)} sections remaining")

        doc = Document()
        self._apply_document_styling(doc)

        reports_dir = output_dir
        os.makedirs(reports_dir, exist_ok=True)

        # Infer comparison/entity context across ALL sections, not the first one that
        # happens to carry entities. Section order is planner-controlled, so a single
        # section that reports is_comparison=False (a synthesis section, say) must not
        # decide the whole report's header just by sorting first.
        is_comparison = any(bool(s.get("is_comparison")) for s in sections)
        entities: list[str] = []
        for s in sections:
            candidate = list(s.get("entities") or [])
            if len(candidate) > len(entities):
                entities = candidate

        # Two or more entities in play is a comparison regardless of what any
        # individual section claimed.
        if len(entities) >= 2:
            is_comparison = True

        rendered_charts: set = set()

        # The planner may emit its own "Executive Summary" / "Conclusion" section
        # (role="synthesize"). Generating the hardcoded ones too would duplicate it.
        need_exec_summary = not self._has_section_titled(sections, "Executive Summary")
        need_conclusion = not self._has_section_titled(sections, "Conclusion")

        from concurrent.futures import ThreadPoolExecutor
        if self.llm:
            with ThreadPoolExecutor(max_workers=3) as summary_executor:
                # Title, summary and conclusion are mutually independent — all three
                # read the finished sections and none reads another's output. The
                # title used to be generated serially before this block, blocking
                # everything behind a full round trip for one line of text.
                title_future = summary_executor.submit(
                    self._generate_report_title, is_comparison, entities, query, sections
                )
                summary_future = summary_executor.submit(
                    self._generate_executive_summary, sections, is_comparison, entities, query
                ) if need_exec_summary else None
                conclusion_future = summary_executor.submit(
                    self._generate_conclusion, sections, is_comparison, entities, query
                ) if need_conclusion else None

                report_title = title_future.result()
                self._add_report_header(doc, report_title, is_comparison, entities)

                if summary_future is not None:
                    doc.add_heading("Executive Summary", level=2)
                    executive_summary = summary_future.result()
                    add_inline_markdown_paragraph(doc, executive_summary)
                    doc.add_paragraph()

                organized_sections = self._organize_sections(sections, query, entities)
                citation_map, references_text = self._build_references_section(organized_sections)

                # Check if this is a multi-vendor comparison
                is_multi_vendor = self._is_multi_vendor_comparison(organized_sections)
                
                for section in organized_sections:
                    if section.get("is_category_header"):
                        doc.add_heading(section.get("heading", "Category"), level=1)
                    else:
                        level = section.get("level", 2)
                        # Skip individual charts for multi-vendor comparisons
                        append_to_doc(doc, section, level=level, citation_map=citation_map, skip_charts=is_multi_vendor, rendered_charts=rendered_charts)
                        doc.add_paragraph()
                
                # Add comprehensive vendor comparison visualization at the end
                if is_multi_vendor:
                    doc.add_heading("Vendor Comparison Summary", level=1)
                    doc.add_paragraph("The following matrix provides a comprehensive visual comparison of all vendors across the evaluated criteria:")
                    
                    viz_path = self._create_vendor_comparison_visualization(organized_sections)
                    if viz_path:
                        doc.add_picture(viz_path, width=Inches(7))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1  # center
                    else:
                        doc.add_paragraph("(Vendor comparison visualization could not be generated)")

                if conclusion_future is not None:
                    doc.add_heading("Conclusion", level=2)
                    conclusion = conclusion_future.result()
                    add_inline_markdown_paragraph(doc, conclusion)

                if references_text:
                    doc.add_paragraph()
                    doc.add_heading("References", level=2)
                    doc.add_paragraph(references_text)
        else:
            # No LLM: the title falls back to its deterministic template.
            report_title = self._generate_report_title(is_comparison, entities, query, sections)
            self._add_report_header(doc, report_title, is_comparison, entities)

            if need_exec_summary:
                doc.add_heading("Executive Summary", level=2)
                executive_summary = self._generate_intro_section(is_comparison, entities)
                doc.add_paragraph(executive_summary)
                doc.add_paragraph()

            citation_map, references_text = self._build_references_section(sections)
            
            # Check if this is a multi-vendor comparison
            is_multi_vendor = self._is_multi_vendor_comparison(sections)
            
            for section in sections:
                # Skip individual charts for multi-vendor comparisons
                append_to_doc(doc, section, level=2, citation_map=citation_map, skip_charts=is_multi_vendor, rendered_charts=rendered_charts)
                doc.add_paragraph()
            
            # Add comprehensive vendor comparison visualization at the end
            if is_multi_vendor:
                doc.add_heading("Vendor Comparison Summary", level=1)
                doc.add_paragraph("The following matrix provides a comprehensive visual comparison of all vendors across the evaluated criteria:")
                
                viz_path = self._create_vendor_comparison_visualization(sections)
                if viz_path:
                    doc.add_picture(viz_path, width=Inches(7))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = 1  # center
                else:
                    doc.add_paragraph("(Vendor comparison visualization could not be generated)")

            if need_conclusion:
                doc.add_heading("Conclusion", level=2)
                conclusion = "This analysis provides insights based on available data from retrieved documents."
                doc.add_paragraph(conclusion)
            if references_text:
                doc.add_paragraph()
                doc.add_heading("References", level=2)
                doc.add_paragraph(references_text)

        now = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"report_{self.model_name}_{now}.docx"
        filepath = os.path.join(reports_dir, filename)
        save_doc(doc, filepath)
        return filepath

    def write_report_typed(
        self,
        sections: List[SectionDraft],
        *,
        query: str = "",
        output_dir: str = "reports",
    ) -> ReportResult:
        """Write the report and return a typed ReportResult model."""
        legacy_sections = [s.to_legacy_dict() for s in sections]
        report_path = self.write_report(legacy_sections, query=query, output_dir=output_dir)
        total_chunks = sum(s.chunks_used for s in sections)
        return ReportResult(
            report_path=report_path,
            sections=sections,
            total_chunks_used=total_chunks,
        )


# Example usage
if __name__ == "__main__":
    doc = Document()
    sample_section = {
        "heading": "Climate Commitments",
        "text": "Both Acme Bank and Globex Bank have committed to net-zero targets...",
        "table": [{"Bank": "Acme Bank", "Target": "Net-zero 2050"},
                  {"Bank": "Globex Bank", "Target": "Net-zero 2050"}],
        "chart_data": {"Acme Bank": 42, "Globex Bank": 36},
        # NEW: tell the pipeline which two entities are being compared
        "entities": ["Acme Bank", "Globex Bank"],
        "is_comparison": True
    }
    agent = ReportWriterAgent(doc)
    agent.write_report([sample_section])
