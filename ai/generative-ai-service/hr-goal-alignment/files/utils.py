# Copyright (c) 2025 Oracle and/or its affiliates.
import streamlit as st
from io import BytesIO
from docx import Document
import oracledb

# -----------------------
# Database Helpers
# -----------------------

def get_employee_id_by_name(connection, name):
    """Fetches employee_id from the database based on name."""
    try:
        cursor = connection.cursor()
        cursor.execute("SELECT employee_id FROM Employees WHERE name = :1", (name,))
        result = cursor.fetchone()
        cursor.close()
        if result:
            return result[0]
    except oracledb.Error as e:
        st.error(f"Error fetching employee ID for {name}: {e}")
    return None

# -----------------------
# Chat Transcript
# -----------------------

def generate_transcript_docx(chat_history, employee_name):
    """Generates a DOCX transcript of the chat."""
    doc = Document()
    doc.add_heading(f"Conversation Transcript with {employee_name}", level=1)

    for exchange in chat_history:
        if "User" in exchange:
            doc.add_paragraph(f"🧑🏽‍💼 You: {exchange['User']}")
        if "Chatbot" in exchange:
            doc.add_paragraph(f"💬 Chatbot: {exchange['Chatbot']}")
            doc.add_paragraph("")  # Spacer

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# -----------------------
# Topic Detection
# -----------------------

def detect_covered_topic(response_text, report_data):
    """Detects which topic was covered in the chatbot response."""
    keywords_map = {
        "quantitative_performance": ["delivery metrics", "improved", "performance metrics", "kpis"],
        "perception_gap": ["perception gap", "rated", "client feedback", "communication skills", "teamwork skills"],
        "manager_context": ["manager", "upcoming meeting", "context with manager", "1:1 discussion"],
        "tenure": ["tenure", "years at company", "length of service", "skill progression"],
        "cross_departmental_contributions": ["cross-department", "collaboration", "led project", "marketing"],
        "career_alignment_prompt": ["career alignment", "long-term", "aspirations", "goals"],
        "key_achievements": ["key achievements", "delivered projects", "task management", "completed project"],
        "skill_development": ["skill development", "leadership skills", "improve in communication", "workshop"],
        "strengths": ["strengths", "delegation", "task prioritization", "analytical skills"],
        "areas_for_improvement": ["areas for improvement", "confidence", "communication", "delegation skills"],
        "alignment_with_career_goals": ["career goals", "leadership abilities", "long-term goals", "senior role"]
    }

    response_lower = response_text.lower()
    for topic, keywords in keywords_map.items():
        if any(keyword in response_lower for keyword in keywords):
            if topic in report_data.get("self_assessment", {}) or topic in report_data.get("manager_briefing", {}):
                return topic
    return None
