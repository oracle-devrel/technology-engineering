# Developed by: Brona Nilsson
"""
DOCX Report Generation Module

Generates professional Word documents with Oracle-branded styling,
tables, and charts. Ported from complex-doc-rag report_writer_agent.py,
simplified for unstructured context passed by the MCP tool.
"""

import os
import re
import json
import math
import uuid
import hashlib
import logging
import datetime
import tempfile
import textwrap

import oci_auth

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# JSON extraction — lightweight replacement for UniversalJSONCleaner
# ---------------------------------------------------------------------------

def extract_json_object(text):
    """Extract and parse a JSON object from LLM output.

    Strips code fences, finds the outermost ``{...}``, attempts parse,
    applies simple fixups, and returns a fallback dict on failure.
    """
    if not text or not isinstance(text, str):
        return {}

    # Strip code fences
    cleaned = re.sub(r"```(?:json)?", "", text).strip()
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].strip()

    # Find outermost { ... }
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return {}

    json_str = cleaned[start : end + 1]

    # First attempt
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        pass

    # Fixup: remove trailing commas before } or ]
    fixed = re.sub(r",\s*([}\]])", r"\1", json_str)
    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # Fixup: replace single quotes with double quotes (crude)
    fixed2 = fixed.replace("'", '"')
    try:
        return json.loads(fixed2)
    except json.JSONDecodeError:
        pass

    logger.warning("extract_json_object: could not parse JSON, returning empty dict")
    return {}


# ---------------------------------------------------------------------------
# DOCX styling helpers (ported from report_writer_agent.py)
# ---------------------------------------------------------------------------

_ORACLE_RED = RGBColor(0xC7, 0x46, 0x34)
_CHARCOAL = RGBColor(0x31, 0x2D, 0x2A)
_MID_GREY = RGBColor(0x74, 0x74, 0x74)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)

_MD_TOKEN_RE = re.compile(r"(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)")


def add_inline_markdown_paragraph(doc, text):
    """Render lightweight inline Markdown (**bold** / *italic*) as a paragraph."""
    p = doc.add_paragraph()
    i = 0
    for m in _MD_TOKEN_RE.finditer(text):
        if m.start() > i:
            p.add_run(text[i : m.start()])
        token = m.group(0)
        if token.startswith("**") or token.startswith("__"):
            run = p.add_run(token[2:-2])
            run.bold = True
        else:
            run = p.add_run(token[1:-1])
            run.italic = True
        i = m.end()
    if i < len(text):
        p.add_run(text[i:])
    return p


def _shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _style_cell(cell, text, *, bold=False, font_size=Pt(8), font_color=_CHARCOAL):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = "Calibri"
    run.font.color.rgb = font_color
    run.bold = bold


def add_table(doc, table_data):
    """Create a Word table with Oracle Red headers and alternating row colors."""
    if not table_data:
        return

    headers = []
    rows_normalized = []

    if isinstance(table_data[0], dict):
        seen = set()
        for row in table_data:
            for k in row.keys():
                if k not in seen:
                    headers.append(k)
                    seen.add(k)
        rows_normalized = table_data
    elif isinstance(table_data[0], (list, tuple)):
        max_len = max(len(row) for row in table_data)
        headers = [f"Col {i+1}" for i in range(max_len)]
        for row in table_data:
            rows_normalized.append(
                {headers[i]: row[i] if i < len(row) else "" for i in range(max_len)}
            )
    else:
        headers = ["Value"]
        rows_normalized = [{"Value": str(row)} for row in table_data]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        _style_cell(cell, str(h), bold=True, font_size=Pt(8.5), font_color=_WHITE)
        _shade_cell(cell, "C74634")

    for row_idx, row in enumerate(rows_normalized):
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            cell = row_cells[i]
            _style_cell(cell, str(row.get(h, "")))
            if row_idx % 2 == 1:
                _shade_cell(cell, "F5F5F5")


def _apply_document_styling(doc):
    """Set Calibri body + Oracle Red headings on the document."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = _CHARCOAL

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = _ORACLE_RED

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = _CHARCOAL


def _add_report_header(doc, title):
    """Add centered title, subtitle, date, and separator."""
    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Analytical Report")
    sub_run.font.size = Pt(12)
    sub_run.font.name = "Calibri"
    sub_run.font.color.rgb = _CHARCOAL
    sub_run.italic = True

    now = datetime.datetime.now()
    date_str = now.strftime("%B %d, %Y")
    time_str = now.strftime("%H:%M")

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"Generated on {date_str} at {time_str}\nPowered by OCI Generative AI"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.name = "Calibri"
    meta_run.font.color.rgb = _MID_GREY

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("\u2500" * 50)
    sep_run.font.color.rgb = _ORACLE_RED
    sep_run.font.size = Pt(8)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Chart generation (ported from report_writer_agent.py)
# ---------------------------------------------------------------------------

_ENTITY_PALETTE = [
    "#C74634",  # Oracle Red
    "#312D2A",  # Charcoal
    "#3A7CA5",  # Steel Blue
    "#747474",  # Mid Grey
    "#D4892A",  # Amber
    "#5B8C5A",  # Sage Green
    "#8E6FAD",  # Purple
    "#2E86AB",  # Ocean Blue
    "#A23B72",  # Magenta
    "#6A7B8B",  # Slate
]


def _parse_numeric(raw):
    """Try to parse a table cell into a float, stripping currency/unit noise."""
    if raw is None:
        return None
    s = str(raw).strip()
    if s.upper() in ("N/A", "", "-", "\u2014", "NONE"):
        return None
    digit_count = sum(1 for c in s if c.isdigit())
    if digit_count == 0:
        return None
    alpha_count = sum(1 for c in s if c.isalpha())
    if alpha_count > digit_count:
        return None
    cleaned = re.sub(r"[^\d.\-eE+]", "", s.replace(",", ""))
    if not cleaned or cleaned in (".", "-", "+"):
        return None
    try:
        val = float(cleaned)
        if math.isnan(val) or math.isinf(val):
            return None
        return val
    except (ValueError, TypeError):
        return None


def detect_units(chart_data, title=""):
    """Detect units of measure from chart data keys and title."""
    currency_patterns = [
        (r"\$|USD|usd|dollar", "USD"),
        (r"\u20ac|EUR|eur|euro", "EUR"),
        (r"\u00a3|GBP|gbp|pound", "GBP"),
    ]
    unit_patterns = [
        (r"million|millions|mn|mln", "Million"),
        (r"billion|billions|bn|bln", "Billion"),
        (r"thousand|thousands", "Thousand"),
        (r"percentage|percent|%", "%"),
        (r"tonnes|tons|tonne|ton", "Tonnes"),
        (r"co2e|CO2e|co2|CO2", "CO2e"),
        (r"employees|headcount|people", "Employees"),
    ]

    combined = title.lower() + " " + " ".join(str(k).lower() for k in chart_data.keys())
    for v in chart_data.values():
        if isinstance(v, str):
            combined += " " + v.lower()

    currency = None
    for pat, unit in currency_patterns:
        if re.search(pat, combined, re.IGNORECASE):
            currency = unit
            break

    scale = None
    for pat, unit in unit_patterns[:3]:
        if re.search(pat, combined, re.IGNORECASE):
            scale = unit
            break

    other_unit = None
    for pat, unit in unit_patterns[3:]:
        if re.search(pat, combined, re.IGNORECASE):
            other_unit = unit
            break

    if currency and scale:
        return f"{scale} {currency}"
    if currency:
        return currency
    if other_unit:
        return other_unit
    if scale:
        return scale

    financial_terms = ["revenue", "sales", "profit", "income", "cost", "expense"]
    if any(t in combined for t in financial_terms):
        return "Million USD"

    return "Value"


def format_value_with_units(value, units):
    """Format a value with appropriate precision based on units."""
    if "%" in units:
        return f"{value:.1f}%"
    if "Million" in units or "Billion" in units:
        return f"{value:,.1f}"
    if value >= 1000:
        return f"{value:,.0f}"
    return f"{value:.1f}"


def _chart_data_from_table(table, entities):
    """Build nested chart_data from a comparison table."""
    chart = {}
    for row in table:
        metric = str(row.get("Metric", ""))
        if not metric or metric == "Unknown Metric":
            continue
        entity_vals = {}
        for ent in entities:
            val = _parse_numeric(row.get(ent))
            if val is not None:
                entity_vals[ent] = val
        if entity_vals:
            chart[metric] = entity_vals
    return chart


def _mine_chart_from_table(table):
    """Generic fallback: extract flat {metric: value} from any table shape."""
    chart = {}
    skip_cols = {"Metric", "Analysis", "Best Value", "Ranking", "analysis", "metric"}
    for row in table:
        if not isinstance(row, dict):
            continue
        metric = str(row.get("Metric", ""))
        if not metric:
            for k, v in row.items():
                if isinstance(v, str) and not _parse_numeric(v):
                    metric = v
                    break
        if not metric or metric == "Unknown Metric":
            continue
        for k, v in row.items():
            if k in skip_cols or str(k) == metric:
                continue
            val = _parse_numeric(v)
            if val is not None:
                chart[metric] = val
                break
    return chart


def _chart_fingerprint(chart_data):
    """Deterministic SHA-256 hex of chart_data for dedup purposes."""

    def _flatten(obj):
        if isinstance(obj, dict):
            return {str(k): _flatten(v) for k, v in sorted(obj.items())}
        if isinstance(obj, (list, tuple)):
            return [_flatten(i) for i in obj]
        return obj

    canonical = json.dumps(_flatten(chart_data), sort_keys=True, default=str)
    return hashlib.sha256(canonical.encode()).hexdigest()


def _make_grouped_chart(grouped_data, title, units, entity_colors, chart_dir):
    """Render a grouped horizontal bar chart. Returns path to saved PNG."""
    metrics = list(grouped_data.keys())
    if not metrics:
        return None

    all_entities = list(entity_colors.keys())
    n_entities = len(all_entities)
    bar_height = 0.8 / n_entities

    fig, ax = plt.subplots(figsize=(12, max(6, len(metrics) * 1.2)))

    for i, ent in enumerate(all_entities):
        positions = [m + i * bar_height for m in range(len(metrics))]
        vals = [grouped_data[metric].get(ent, 0) for metric in metrics]
        bars = ax.barh(positions, vals, height=bar_height,
                       label=ent, color=entity_colors[ent])
        for bar in bars:
            width = bar.get_width()
            if width != 0:
                formatted = format_value_with_units(width, units)
                ax.annotate(formatted,
                            xy=(width, bar.get_y() + bar.get_height() / 2),
                            xytext=(5, 0), textcoords="offset points",
                            ha="left", va="center", fontsize=8)

    center_offsets = [m + bar_height * (n_entities - 1) / 2 for m in range(len(metrics))]
    wrapped = ["\n".join(textwrap.wrap(m, width=35)) for m in metrics]
    ax.set_yticks(center_offsets)
    ax.set_yticklabels(wrapped)
    ax.set_xlabel(units)
    ax.set_title(title[:100])
    ax.legend(loc="lower right", framealpha=0.9)
    ax.grid(axis="x", linestyle="--", alpha=0.6)
    fig.tight_layout()

    filename = f"chart_{uuid.uuid4().hex}.png"
    path = os.path.join(chart_dir, filename)
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart(chart_data, title="", chart_dir="/tmp"):
    """Generate a bar chart (flat or grouped). Returns path to PNG or None."""
    if not chart_data:
        return None

    os.makedirs(chart_dir, exist_ok=True)

    units = detect_units(chart_data, title)
    if units and units != "Value" and units.lower() not in title.lower():
        title = f"{title} ({units})"

    # --- Grouped mode ---
    grouped = {}
    for k, v in chart_data.items():
        if isinstance(v, dict):
            nums = {}
            for sub_k, sub_v in v.items():
                try:
                    n = float(sub_v)
                    if not math.isnan(n) and not math.isinf(n):
                        nums[str(sub_k)] = n
                except Exception:
                    continue
            if nums:
                grouped[str(k)[:80]] = nums

    if grouped:
        seen_ents = []
        for sub in grouped.values():
            for e in sub:
                if e not in seen_ents:
                    seen_ents.append(e)
        entity_colors = {}
        for i, e in enumerate(seen_ents):
            entity_colors[e] = _ENTITY_PALETTE[i % len(_ENTITY_PALETTE)]
        return _make_grouped_chart(grouped, title, units, entity_colors, chart_dir)

    # --- Flat mode ---
    clean = {}
    for k, v in chart_data.items():
        if isinstance(v, list):
            if all(isinstance(i, (int, float)) for i in v):
                v = v[-1]
            else:
                continue
        try:
            num = float(v)
            if not math.isnan(num) and not math.isinf(num):
                clean[str(k)[:80]] = num
        except Exception:
            continue

    if not clean:
        return None

    labels = list(clean.keys())
    values = list(clean.values())

    max_label_length = max(len(l) for l in labels) if labels else 0
    horizontal = len(clean) > 6 or max_label_length > 30

    fig, ax = plt.subplots(figsize=(12, 8))

    if horizontal:
        wrapped_labels = ["\n".join(textwrap.wrap(l, width=40)) for l in labels]
        bars = ax.barh(wrapped_labels, values, color=_ENTITY_PALETTE[0])
        ax.set_xlabel(units)
        for bar in bars:
            width = bar.get_width()
            formatted = format_value_with_units(width, units)
            ax.annotate(formatted,
                        xy=(width, bar.get_y() + bar.get_height() / 2),
                        xytext=(5, 0), textcoords="offset points",
                        ha="left", va="center", fontsize=8)
    else:
        wrapped_labels = ["\n".join(textwrap.wrap(l, width=15)) for l in labels]
        bars = ax.bar(range(len(labels)), values, color=_ENTITY_PALETTE[0])
        ax.set_ylabel(units)
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(wrapped_labels, ha="center", va="top")
        for bar in bars:
            height = bar.get_height()
            formatted = format_value_with_units(height, units)
            ax.annotate(formatted,
                        xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 5), textcoords="offset points",
                        ha="center", va="bottom", fontsize=8)

    ax.set_title(title[:100])
    ax.grid(axis="y" if not horizontal else "x", linestyle="--", alpha=0.6)
    fig.tight_layout()

    filename = f"chart_{uuid.uuid4().hex}.png"
    path = os.path.join(chart_dir, filename)
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# Section assembly
# ---------------------------------------------------------------------------

def append_to_doc(doc, section_data, chart_dir, rendered_charts):
    """Append a section (heading + text + table + chart) to the document."""
    heading = section_data.get("heading", "Untitled Section")
    doc.add_heading(heading, level=2)

    text = section_data.get("text", "").strip()
    if text:
        add_inline_markdown_paragraph(doc, text)

    table_data = section_data.get("table", [])
    if isinstance(table_data, dict):
        table_data = [table_data]
    if isinstance(table_data, list) and table_data:
        add_table(doc, table_data)

    chart_data = section_data.get("chart_data", {})
    if isinstance(chart_data, dict) and chart_data:
        fp = _chart_fingerprint(chart_data)
        if fp not in rendered_charts:
            rendered_charts.add(fp)
            chart_path = make_chart(chart_data, title=heading, chart_dir=chart_dir)
            if chart_path:
                doc.add_picture(chart_path, width=Inches(6))
                doc.paragraphs[-1].alignment = 1  # center


# ---------------------------------------------------------------------------
# Section writing (adapted for unstructured context)
# ---------------------------------------------------------------------------

def _write_section_core(topic, context_trimmed):
    """Shared LLM call + JSON normalization for writing a report section.

    Called by both write_section() (KB path) and write_section_with_context()
    (orchestrated path where Agent Hub provides context).

    Returns a dict with keys: heading, text, table, chart_data.
    """
    prompt = f"""You are writing one section of a professional analytical report.

SECTION TOPIC: {topic}

SOURCE DATA:
{context_trimmed}

Write this section and return ONLY a JSON object with these keys:
- "heading": A descriptive heading for this section (string)
- "text": 2-4 paragraphs of analytical text (string). Be specific and data-driven.
- "table": A list of row objects for a comparison/data table. Each row is a dict.
  For comparisons use the ACTUAL entity/period names as column headers, e.g.:
  {{"Metric": "Revenue", "FY2025": "$12.3B", "FY2026": "$14.1B", "Analysis": "14.6% growth"}}
  {{"Metric": "Revenue", "Acme Corp": "$5.2B", "Beta Inc": "$3.8B", "Analysis": "Acme leads"}}
  For single-entity use: {{"Metric": "...", "Value": "..."}}
  IMPORTANT: Never use generic names like "Entity1" or "Entity2" — use the real names from the data.
  Include 3-8 rows with specific data from the source material.
- "chart_data": A dict of metric names to numeric values for charting.
  For comparisons use nested format with REAL entity names: {{"Revenue": {{"FY2025": 12.3, "FY2026": 14.1}}, ...}}
  For single-entity use flat format: {{"Metric1": 123, "Metric2": 456, ...}}
  ONLY numeric values - no text, no units, no currency symbols.

RULES:
1. Use ONLY data from the source material. Do not fabricate numbers.
2. If data is missing, note it in the text and use "N/A" in tables.
3. Never use possessive apostrophes (write "Oracle revenue" not "Oracle's revenue").
4. Return valid JSON only.
5. Include at least 3 metrics in chart_data if numeric data is available."""

    try:
        response = oci_auth.llm_chat(prompt, max_tokens=4000)
        parsed = extract_json_object(response)

        if not parsed.get("heading"):
            parsed["heading"] = topic
        if not parsed.get("text"):
            parsed["text"] = ""

        # Normalize table
        table = parsed.get("table", [])
        if isinstance(table, dict):
            table = [table]
        if not isinstance(table, list):
            table = []
        parsed["table"] = table

        # Normalize chart_data
        chart_data = parsed.get("chart_data", {})
        if not isinstance(chart_data, dict):
            chart_data = {}

        # If chart_data is empty, try mining from table
        has_numeric = any(
            _parse_numeric(v) is not None
            for v in chart_data.values()
            if not isinstance(v, dict)
        ) if chart_data else False

        has_nested = any(isinstance(v, dict) for v in chart_data.values()) if chart_data else False

        if not has_numeric and not has_nested and table:
            mined = _mine_chart_from_table(table)
            if mined:
                chart_data = mined

        parsed["chart_data"] = chart_data
        return parsed

    except Exception as e:
        logger.error("Failed to write section '%s': %s", topic, e)
        return {
            "heading": topic,
            "text": "Could not generate this section due to an error.",
            "table": [],
            "chart_data": {},
        }


def write_section(topic, query, initial_context, chart_dir=None):
    """Write a single report section via one LLM call.

    Searches the KB for targeted context on this section's topic,
    then combines it with the initial context for richer data.

    Returns a dict with keys: heading, text, table, chart_data.
    """
    # Search the KB specifically for this section's topic
    search_query = f"{query} {topic}"
    logger.info("write_section: searching KB for '%s'", search_query)
    kb_result = oci_auth.kb_search(search_query)

    # Combine KB results with initial context, KB results first (more relevant)
    combined_parts = []
    if kb_result:
        combined_parts.append(f"=== KB Search Results for '{topic}' ===\n{kb_result}")
    if initial_context:
        combined_parts.append(f"=== Initial Context ===\n{initial_context}")
    combined_context = "\n\n".join(combined_parts)

    context_trimmed = combined_context[:8000] if len(combined_context) > 8000 else combined_context
    return _write_section_core(topic, context_trimmed)


def write_section_with_context(topic, query, context):
    """Write a single report section using pre-gathered context (no KB search).

    Used in the orchestrated path where Agent Hub provides the context
    via its own RAG tool.

    Returns a dict with keys: heading, text, table, chart_data.
    """
    context_trimmed = context[:8000] if len(context) > 8000 else context
    return _write_section_core(topic, context_trimmed)


# ---------------------------------------------------------------------------
# Report orchestrator
# ---------------------------------------------------------------------------

def _plan_sections(query, context):
    """Ask the LLM to plan 3-6 section topics from the query and context."""
    context_preview = context[:4000] if len(context) > 4000 else context

    prompt = f"""You are planning the structure of an analytical report.

USER REQUEST: {query}

AVAILABLE DATA (preview):
{context_preview}

Based on the request and data, list 3-6 section topics for this report.
Each topic should be a short phrase (3-8 words) describing what that section covers.

Return ONLY a JSON object:
{{"sections": ["Topic 1", "Topic 2", "Topic 3", ...]}}

Guidelines:
- Cover the most important aspects of the data
- If comparing entities, include sections for each key comparison dimension
- Include sections for key metrics, trends, and notable findings
- Do NOT include "Executive Summary" or "Conclusion" — those are added automatically"""

    try:
        response = oci_auth.llm_chat(prompt, max_tokens=500)
        parsed = extract_json_object(response)
        sections = parsed.get("sections", [])
        if isinstance(sections, list) and len(sections) >= 2:
            return sections[:6]
    except Exception as e:
        logger.warning("Section planning failed: %s", e)

    # Fallback: generic sections
    return [
        "Key Metrics and Data Overview",
        "Detailed Analysis",
        "Trends and Notable Findings",
    ]


def plan_sections(query, context):
    """Public wrapper around _plan_sections for use by the MCP server."""
    return _plan_sections(query, context)


def assemble_docx(title, query, sections):
    """Assemble pre-written sections into a final DOCX report.

    Generates the executive summary and conclusion via LLM calls,
    renders charts, and returns the path to the saved DOCX file.

    Used by both build_docx_report() and the orchestrated assemble_report
    MCP tool.
    """
    logger.info("assemble_docx: starting for '%s' with %d sections", title, len(sections))

    chart_dir = tempfile.mkdtemp(prefix="docx_charts_")

    # Generate executive summary
    section_summaries = []
    for s in sections:
        h = s.get("heading", "")
        t = s.get("text", "")
        if t:
            section_summaries.append(f"{h}: {t[:300]}")

    summaries_text = "\n\n".join(section_summaries)

    # Fast mode: use templates for exec summary and conclusion (saves ~8s of LLM calls)
    exec_summary = (
        f"This report provides a detailed analysis addressing the following request: {query}. "
        f"The analysis is organized into {len(sections)} sections, each examining a different "
        f"aspect of the available data. Key findings from each section are summarized below."
    )
    for s in section_summaries[:3]:
        exec_summary += f"\n\n{s}"

    conclusion = (
        f"This analysis has examined {len(sections)} key dimensions of the requested topic. "
        "The findings presented are based on data extracted from the source documents in the knowledge base. "
        "For further detail on any specific area, additional targeted analysis may be conducted."
    )

    # Assemble DOCX
    doc = Document()
    _apply_document_styling(doc)
    _add_report_header(doc, title)

    # Executive summary
    doc.add_heading("Executive Summary", level=2)
    add_inline_markdown_paragraph(doc, exec_summary)
    doc.add_paragraph()

    # Sections
    rendered_charts = set()
    for section in sections:
        append_to_doc(doc, section, chart_dir, rendered_charts)
        doc.add_paragraph()

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    add_inline_markdown_paragraph(doc, conclusion)

    # Save to temp file
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w\s-]", "", title)[:50].strip().replace(" ", "_")
    filename = f"report_{safe_title}_{timestamp}.docx"
    filepath = os.path.join(tempfile.gettempdir(), filename)
    doc.save(filepath)

    logger.info("assemble_docx: saved to %s", filepath)

    # Clean up chart PNGs
    try:
        for f in os.listdir(chart_dir):
            os.remove(os.path.join(chart_dir, f))
        os.rmdir(chart_dir)
    except Exception:
        pass

    return filepath


def build_docx_report(title, query, context):
    """Build a complete DOCX report. Returns path to the generated file.

    Optimized for speed (~25s) to fit within OCI Agent tool timeout:
    1. Plan section topics
    2. Batch KB searches in parallel
    3. Write all sections in parallel using KB results
    4. Assemble DOCX (exec summary + conclusion in parallel)
    """
    import concurrent.futures

    logger.info("build_docx_report: starting for '%s'", title)

    # 1. Plan sections (up to 5 when context provided, 3 with KB search)
    max_sections = 5 if (context and len(context.strip()) > 50) else 3
    section_topics = _plan_sections(query, context)[:max_sections]
    logger.info("build_docx_report: planned %d sections: %s", len(section_topics), section_topics)

    # 2. Get context for each section
    has_context = bool(context and len(context.strip()) > 50)

    if has_context:
        # Agent provided RAG results — skip KB search (saves ~10s)
        logger.info("build_docx_report: using provided context (%d chars), skipping KB search", len(context))
        kb_results = [None] * len(section_topics)
    else:
        # No context — batch KB searches in parallel
        def _kb_search(topic):
            search_query = f"{query} {topic}"
            logger.info("build_docx_report: KB search for '%s'", topic)
            return oci_auth.kb_search(search_query)

        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            kb_results = list(executor.map(_kb_search, section_topics))

    # 3. Write all sections in parallel
    def _write(args):
        topic, kb_result = args
        combined = ""
        if kb_result:
            combined += f"=== KB Results for '{topic}' ===\n{kb_result}\n\n"
        if context:
            combined += f"=== Additional Context ===\n{context}\n\n"
        combined = combined[:8000] if len(combined) > 8000 else combined
        logger.info("build_docx_report: writing section '%s'", topic)
        return _write_section_core(topic, combined)

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        sections = list(executor.map(_write, zip(section_topics, kb_results)))

    # 4. Assemble into final DOCX
    return assemble_docx(title, query, sections)
